"""Generate .docx architecture documents for AI Debate Platform."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, hex_color="1F497D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_style(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)


def add_header_row(table, headers, hex_color="1F497D"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, hex_color)


def make_table(doc, headers, rows_data, hex_color="1F497D"):
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    set_table_style(tbl)
    add_header_row(tbl, headers, hex_color)
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row if isinstance(row, (list, tuple)) else [row]):
            tbl.rows[i + 1].cells[j].text = val
    doc.add_paragraph()
    return tbl


def code_block(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────────
# 1. 시스템 아키텍처 설계서
# ──────────────────────────────────────────────────────────────────
def build_system_arch():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Title page
    t = doc.add_heading("시스템 아키텍처 설계서", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("작성일: 2026-03-09    버전: v1.0    AI 에이전트 토론 플랫폼").italic = True
    doc.add_page_break()

    # TOC
    doc.add_heading("목차", 1)
    for item in [
        "1. 서비스 개요",
        "2. 인프라 구성",
        "3. 전체 시스템 아키텍처",
        "4. 백엔드 구조",
        "5. 프론트엔드 구조",
        "6. 데이터베이스 구조",
        "7. 토론 엔진 흐름",
        "8. API 엔드포인트 명세",
        "9. 사용자 역할 및 접근 제어",
        "10. 실시간 스트리밍 구조",
        "11. 관측성 및 모니터링",
        "12. 성능 목표",
        "13. 배포 구성",
    ]:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # 1. 서비스 개요
    doc.add_heading("1. 서비스 개요", 1)
    doc.add_paragraph(
        "AI 에이전트 토론 플랫폼은 AI 에이전트끼리 실시간으로 토론을 벌이고, "
        "사용자가 관전·예측투표·시즌 랭킹을 즐기는 플랫폼이다."
    )
    make_table(
        doc,
        ["항목", "내용"],
        [
            ("서비스 단계", "프로토타입 (동시 접속 10명 이하)"),
            ("월 예상 비용", "~$130 (EC2 ~$15 + RunPod ~$114) + LLM API 비용 (사용량 비례)"),
            ("주요 특징", "AI 에이전트 토론, 실시간 관전, ELO 랭킹, 예측투표, 토너먼트, 턴 검토 시스템"),
        ],
    )
    doc.add_heading("핵심 기능", 2)
    for f in [
        "AI 에이전트 토론: 사용자가 에이전트(성격·모델·프롬프트)를 직접 생성하고 토론에 참가",
        "실시간 관전: Redis Pub/Sub + SSE로 토론 진행 상황 실시간 브로드캐스트",
        "턴 검토 시스템: gpt-5-nano가 매 발언을 검토 (논리 오류·허위 주장·주제 이탈 탐지, 벌점 부여)",
        "ELO 랭킹·시즌: ELO 기반 시즌 랭킹, 승급전(3판 2선승)/강등전(1판) 자동 생성",
        "예측투표: 매치 시작 전 사용자 승자 예측, 완료 후 결과 공개",
        "토너먼트: 대진표 자동 생성, 단계별 진행",
        "LLM 모델 전환: 에이전트별 LLM 모델 선택 가능 (OpenAI/Anthropic/Google/RunPod/Local)",
        "토큰 사용량 추적: 사용자별 LLM 토큰 사용량 실시간 추적 및 비용 산출",
        "관리자 대시보드: 매치 관리, 시즌/토너먼트 관리, 모니터링, 사용량/과금 현황",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    # 2. 인프라 구성
    doc.add_heading("2. 인프라 구성", 1)
    make_table(
        doc,
        ["역할", "기술"],
        [
            ("Frontend", "Next.js 15 + React 19 + Zustand, Tailwind CSS"),
            ("Backend", "Python 3.12 + FastAPI, SQLAlchemy 2.0 async"),
            ("Database", "PostgreSQL 16 (Docker) — 18개 테이블"),
            ("Cache/Pub-Sub", "Redis (Docker)"),
            ("LLM Inference", "RunPod Serverless + SGLang (기본) + OpenAI/Anthropic/Google API"),
            ("Streaming", "SSE (Server-Sent Events)"),
            ("Observability", "Langfuse + Prometheus + Grafana + Sentry"),
            ("Infra", "AWS EC2 t4g.small (서울 ap-northeast-2) + RunPod Serverless (미국)"),
            ("Container", "Docker Compose"),
        ],
    )
    doc.add_heading("인프라 배치", 2)
    for d in [
        "EC2 t4g.small (서울): FastAPI 백엔드, PostgreSQL, Redis를 Docker Compose로 함께 운용",
        "RunPod Serverless (미국): Llama 3.3 70B SGLang 추론 서버, 콜드 스타트 ~15초",
        "외부 LLM API: OpenAI, Anthropic, Google — RTT ~150ms (에이전트 발언 생성 시)",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    # 3. 전체 시스템 아키텍처
    doc.add_heading("3. 전체 시스템 아키텍처", 1)
    doc.add_paragraph(
        "아래 다이어그램은 사용자 브라우저부터 EC2 백엔드, 외부 LLM API, 관측성 도구까지의 "
        "전체 데이터 흐름을 나타낸다."
    )
    code_block(
        doc,
        (
            "┌─ 사용자 브라우저 ─────────────────────────────────┐\n"
            "│  Next.js 15 Frontend  (Zustand + SSE Client)    │\n"
            "└──────────────────┬────────────────────────────────┘\n"
            "                   │ HTTPS / SSE / WebSocket\n"
            "┌──────────────────▼────────────────────────────────┐\n"
            "│  EC2 서울 t4g.small                               │\n"
            "│  ┌─ FastAPI ───────────────────────────────────┐  │\n"
            "│  │  /api/auth   /api/agents   /api/matches      │  │\n"
            "│  │  /api/topics /api/tournaments /api/ws         │  │\n"
            "│  │  /api/models /api/usage   /api/admin          │  │\n"
            "│  └─────────────────────────────────────────────┘  │\n"
            "│  PostgreSQL 16 (18 tables)  Redis Pub/Sub + Cache │\n"
            "└──────────────────┬────────────────────────────────┘\n"
            "                   │ LLM HTTP (~150ms RTT)\n"
            "┌──────────────────▼────────────────────────────────┐\n"
            "│  LLM 프로바이더                                   │\n"
            "│  OpenAI API  /  Anthropic API  /  Google API      │\n"
            "│  RunPod Serverless (Llama 3.3 70B / SGLang)       │\n"
            "└───────────────────────────────────────────────────┘"
        ),
    )
    p = doc.add_paragraph()
    p.add_run("FigJam 다이어그램: ").bold = True
    p.add_run(
        "https://www.figma.com/online-whiteboard/create-diagram/ccd7315e-f8c2-44a1-813f-1b6423ceca7a"
    )

    # 4. 백엔드 구조
    doc.add_heading("4. 백엔드 구조", 1)
    doc.add_heading("디렉토리 구조", 2)
    code_block(
        doc,
        (
            "backend/app/\n"
            "├── main.py                  # FastAPI 앱, 라우터 등록\n"
            "├── api/                     # 라우터 레이어 (입력 검증 + HTTP 응답만)\n"
            "│   ├── auth.py\n"
            "│   ├── debate_agents.py\n"
            "│   ├── debate_matches.py\n"
            "│   ├── debate_topics.py\n"
            "│   ├── debate_tournaments.py\n"
            "│   ├── debate_ws.py         # WebSocket (로컬 에이전트 연결)\n"
            "│   └── admin/\n"
            "│       ├── debate/          # 매치·시즌·토너먼트·에이전트 관리\n"
            "│       └── system/          # 사용자·LLM 모델·모니터링·사용량 관리\n"
            "├── core/                    # 인프라 설정\n"
            "├── models/                  # SQLAlchemy ORM 모델 (18개)\n"
            "├── schemas/                 # Pydantic v2 입출력 스키마\n"
            "└── services/                # 비즈니스 로직"
        ),
    )
    doc.add_heading("주요 서비스 목록", 2)
    make_table(
        doc,
        ["파일", "역할"],
        [
            ("debate_agent_service.py", "에이전트 CRUD, 랭킹, 갤러리, 클론, H2H, 버전 관리"),
            ("debate_match_service.py", "매치 조회, 하이라이트, 요약 리포트 생성"),
            ("debate_matching_service.py", "큐 등록/취소, 자동 매칭(DebateAutoMatcher), ready_up"),
            ("debate_engine.py", "토론 실행 루프 (턴 실행 → 검토 → 판정 → 결과 저장)"),
            ("debate_orchestrator.py", "LLM 검토 + 최적화 병렬 실행(OptimizedDebateOrchestrator)"),
            ("debate_broadcast.py", "SSE 이벤트 발행/구독, 관전자 수 관리"),
            ("debate_ws_manager.py", "WebSocket 연결 관리 (로컬 에이전트 인증·메시지 라우팅)"),
            ("inference_client.py", "LLM 호출 단일 진입점 (Langfuse 추적, 토큰 로깅, provider 분기)"),
            ("debate_season_service.py", "시즌 생성/종료, 시즌 ELO 집계, 보상 지급"),
            ("debate_promotion_service.py", "승급전/강등전 시리즈 생성·진행·완료 처리"),
            ("debate_tournament_service.py", "토너먼트 대진표 생성·진행"),
        ],
    )
    doc.add_heading("레이어 책임 원칙", 2)
    make_table(
        doc,
        ["레이어", "책임", "금지 사항"],
        [
            ("api/ (라우터)", "HTTP 요청 수신, 입력 검증, 응답 포맷", "DB 직접 쿼리, 비즈니스 로직"),
            ("services/", "비즈니스 로직, DB 조작", "HTTP 관련 코드"),
            ("models/", "ORM 정의, 테이블 구조", "비즈니스 로직"),
            ("schemas/", "입출력 데이터 검증", "DB 접근"),
            ("core/", "인프라 설정, 공통 의존성", "도메인 로직"),
        ],
    )

    # 5. 프론트엔드 구조
    doc.add_heading("5. 프론트엔드 구조", 1)
    doc.add_heading("디렉토리 구조", 2)
    code_block(
        doc,
        (
            "frontend/src/\n"
            "├── app/\n"
            "│   ├── (user)/              # 사용자 라우트 그룹\n"
            "│   │   ├── debate/          # 토론 목록, 매치 관전, 갤러리, 랭킹\n"
            "│   │   ├── agents/          # 에이전트 생성/편집\n"
            "│   │   ├── seasons/         # 시즌 랭킹\n"
            "│   │   ├── tournaments/     # 토너먼트 대진표\n"
            "│   │   └── mypage/          # 내 정보, 사용량\n"
            "│   ├── admin/               # 관리자 라우트 그룹\n"
            "│   └── api/[...path]/       # Next.js → FastAPI SSE 프록시\n"
            "├── components/\n"
            "│   ├── debate/              # DebateViewer, TurnBubble, PromotionBadge 등\n"
            "│   ├── admin/               # 관리자 UI 컴포넌트\n"
            "│   └── layout/              # 공통 레이아웃\n"
            "├── stores/                  # Zustand 상태 관리\n"
            "└── lib/\n"
            "    ├── api.ts               # API 호출 유틸리티\n"
            "    ├── auth.ts              # 인증 처리\n"
            "    └── agentWebSocket.ts    # 로컬 에이전트 WebSocket 클라이언트"
        ),
    )
    doc.add_heading("주요 Zustand 상태 (debateStore.ts)", 2)
    make_table(
        doc,
        ["상태", "설명"],
        [
            ("currentMatch", "현재 관전 중인 매치 정보"),
            ("turns", "누적된 턴 발언 목록"),
            ("turnReviews", "턴별 LLM 검토 결과 (벌점, 논리 점수)"),
            ("prediction", "사용자의 예측투표 상태"),
            ("viewerCount", "현재 관전자 수"),
            ("replayMode / replayIndex", "리플레이 모드 상태"),
        ],
    )

    # 6. 데이터베이스 구조
    doc.add_heading("6. 데이터베이스 구조", 1)
    doc.add_heading("테이블 목록 (18개)", 2)
    make_table(
        doc,
        ["모델", "테이블", "설명"],
        [
            ("User", "users", "사용자 계정, 역할(user/admin/superadmin), 크레딧 잔액"),
            ("LLMModel", "llm_models", "등록된 LLM 모델 (provider, 비용, 활성화 여부)"),
            ("TokenUsageLog", "token_usage_logs", "LLM 호출 토큰·비용 기록"),
            ("DebateAgent", "debate_agents", "에이전트 (소유자, provider, ELO, 공개 여부, 승급전 상태)"),
            ("DebateAgentVersion", "debate_agent_versions", "에이전트 버전 이력 (system_prompt 스냅샷)"),
            ("DebateAgentSeasonStats", "debate_agent_season_stats", "시즌별 ELO·전적 분리 집계"),
            ("DebateAgentTemplate", "debate_agent_templates", "관리자 제공 에이전트 템플릿"),
            ("DebateTopic", "debate_topics", "토론 주제 (등록자, 승인 상태)"),
            ("DebateMatch", "debate_matches", "매치 (참가자, 형식, 상태, 결과, 시즌/시리즈 연결)"),
            ("DebateMatchParticipant", "debate_match_participants", "멀티에이전트 매치 참가자 목록"),
            ("DebateMatchPrediction", "debate_match_predictions", "사용자 예측투표"),
            ("DebateMatchQueue", "debate_match_queues", "매칭 대기 큐"),
            ("DebateTurnLog", "debate_turn_logs", "턴별 발언·검토 결과·점수 기록"),
            ("DebatePromotionSeries", "debate_promotion_series", "승급전/강등전 시리즈 상태"),
            ("DebateSeason", "debate_seasons", "시즌 기간·상태"),
            ("DebateSeasonResult", "debate_season_results", "시즌 종료 시 최종 순위 스냅샷"),
            ("DebateTournament", "debate_tournaments", "토너먼트 대진표·상태"),
            ("DebateTournamentEntry", "debate_tournament_entries", "토너먼트 참가 에이전트 목록"),
        ],
    )
    doc.add_heading("DB 컨벤션", 2)
    for c in [
        "테이블/컬럼: snake_case",
        "PK: id (UUID, gen_random_uuid())",
        "FK: {참조테이블_단수}_id",
        "인덱스: idx_{테이블}_{컬럼}",
        "타임스탬프: TIMESTAMPTZ (타임존 포함)",
        "Enum 대신 CHECK 제약조건 사용",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    # 7. 토론 엔진 흐름
    doc.add_heading("7. 토론 엔진 흐름", 1)
    doc.add_heading("전체 흐름", 2)
    code_block(
        doc,
        (
            "큐 등록 → DebateAutoMatcher 감지 → ready_up() → DebateMatch 생성\n"
            "    → debate_engine.run_match()\n"
            "        ├─ 턴 루프 (N 라운드)\n"
            "        │   ├─ 에이전트 발언 생성 (LLM 호출 or WebSocket)\n"
            "        │   └─ OptimizedDebateOrchestrator.review_turn()\n"
            "        │       └─ asyncio.gather(A 검토, B 실행) 병렬 실행\n"
            "        └─ judge() → 최종 판정 → ELO 갱신 → 승급전 체크\n"
            "    → SSE 이벤트 발행 (debate_broadcast → Redis Pub/Sub)"
        ),
    )
    doc.add_heading("OptimizedOrchestrator 병렬 처리", 2)
    doc.add_paragraph(
        "턴 루프에서 A 에이전트 발언 검토와 B 에이전트 발언 생성을 asyncio.gather()로 동시 실행한다. "
        "순차 처리 대비 시간을 37% 단축한다."
    )
    code_block(
        doc,
        (
            "# asyncio.gather로 A 검토와 B 실행을 병렬화 — 턴 지연 37% 단축\n"
            "review_a, result_b = await asyncio.gather(\n"
            "    orchestrator.review_turn(turn_a),\n"
            "    engine.execute_turn(agent_b)\n"
            ")"
        ),
    )
    p = doc.add_paragraph()
    p.add_run("FigJam 시퀀스 다이어그램: ").bold = True
    p.add_run(
        "https://www.figma.com/online-whiteboard/create-diagram/fdf2f969-2ca9-4695-8ca9-6ed628c857d0"
    )
    doc.add_heading("승급전/강등전 시스템", 2)
    make_table(
        doc,
        ["종류", "조건", "규칙"],
        [
            ("승급전", "ELO 일정 임계값 초과", "3판 2선승 (required_wins=2)"),
            ("강등전", "ELO 일정 임계값 미만", "1판 필승 (required_wins=1)"),
        ],
    )

    # 8. API 엔드포인트
    doc.add_heading("8. API 엔드포인트 명세", 1)
    make_table(
        doc,
        ["경로", "파일", "설명"],
        [
            ("GET /health", "health.py", "서버 상태 확인"),
            ("/api/auth/*", "auth.py", "회원가입, 로그인, 토큰 갱신"),
            ("/api/agents/*", "debate_agents.py", "에이전트 CRUD, 랭킹, 갤러리, H2H"),
            ("/api/topics/*", "debate_topics.py", "토픽 등록/조회/매칭 큐"),
            ("/api/matches/*", "debate_matches.py", "매치 조회, SSE 스트리밍, 예측투표, 요약"),
            ("/api/tournaments/*", "debate_tournaments.py", "토너먼트 CRUD, 대진표"),
            ("/api/models/*", "models.py", "LLM 모델 목록 조회, 선호 모델 설정"),
            ("/api/usage/*", "usage.py", "내 토큰 사용량 조회"),
            ("/api/ws/debate/*", "debate_ws.py", "WebSocket (로컬 에이전트 전용)"),
            ("/api/admin/users/*", "admin/system/users.py", "사용자 조회/역할 변경"),
            ("/api/admin/models/*", "admin/system/llm_models.py", "LLM 모델 등록/수정/활성화"),
            ("/api/admin/usage/*", "admin/system/usage.py", "전체 사용량 현황"),
            ("/api/admin/monitoring/*", "admin/system/monitoring.py", "토큰/비용 모니터링"),
            ("/api/admin/debate/*", "admin/debate/", "매치 강제실행, 시즌/토너먼트 관리"),
        ],
    )
    doc.add_heading("인증 방식", 2)
    for a in [
        'REST API: Authorization: Bearer <JWT> 헤더',
        'WebSocket: 연결 후 첫 메시지로 인증 ({"type": "auth", "token": "<JWT>"})',
        "인증 실패 또는 5초 내 미전송 시 연결 즉시 종료",
    ]:
        doc.add_paragraph(a, style="List Bullet")
    doc.add_heading("SSE 스트리밍 이벤트", 2)
    make_table(
        doc,
        ["이벤트 타입", "설명"],
        [
            ("match_start", "토론 시작 알림"),
            ("turn", "에이전트 발언"),
            ("turn_review", "턴 검토 결과 (벌점, 논리 점수)"),
            ("match_end", "토론 완료, 최종 결과"),
            ("series_update", "승급전/강등전 시리즈 상태 변경"),
            ("viewer_count", "현재 관전자 수 갱신"),
        ],
    )

    # 9. RBAC
    doc.add_heading("9. 사용자 역할 및 접근 제어", 1)
    make_table(
        doc,
        ["역할", "접근 범위", "주요 기능"],
        [
            (
                "user",
                "토론 관전, 에이전트 생성/편집, 예측투표, 랭킹 조회, 사용량 조회",
                "에이전트 커스터마이징, 큐 등록, 토너먼트 참가, LLM 모델 선택",
            ),
            (
                "admin",
                "관리자 대시보드 + 사용자 화면 전체 (읽기 위주)",
                "매치 관리, 시즌/토너먼트 관리, 모니터링, 에이전트 모더레이션",
            ),
            (
                "superadmin",
                "admin 전체 + 파괴적 작업",
                "사용자 삭제/역할 변경, LLM 모델 등록/수정, 시스템 설정, 쿼터 관리",
            ),
        ],
    )
    doc.add_heading("접근 제어 원칙", 2)
    for r in [
        "관리자 API: Depends(require_admin) 또는 Depends(require_superadmin) 필수",
        "파괴적 작업(삭제/역할 변경): require_superadmin 필수",
        "사용자는 자신의 리소스만 접근 가능 (소유권 체크 필수)",
        "소유권 실패 → HTTP 403, 미존재 → HTTP 404",
    ]:
        doc.add_paragraph(r, style="List Bullet")

    # 10. 실시간 스트리밍
    doc.add_heading("10. 실시간 스트리밍 구조", 1)
    doc.add_heading("Redis Pub/Sub + SSE 흐름", 2)
    code_block(
        doc,
        (
            'DebateEngine\n'
            '    → Redis PUBLISH "debate:{match_id}"\n'
            '        → FastAPI SSE 구독자\n'
            '            → HTTP/2 SSE 스트림\n'
            '                → 브라우저 EventSource'
        ),
    )
    for n in [
        "Redis 채널 이름: debate:{match_id}",
        "Next.js API Route (/app/api/[...path]/route.ts)가 SSE 프록시 역할 담당",
        "관전자 수: Redis INCR/DECR로 실시간 집계",
    ]:
        doc.add_paragraph(n, style="List Bullet")

    # 11. 관측성
    doc.add_heading("11. 관측성 및 모니터링", 1)
    make_table(
        doc,
        ["도구", "역할"],
        [
            ("Langfuse", "LLM 호출 추적 (모델, 토큰, 응답 시간, 비용)"),
            ("Prometheus", "애플리케이션 메트릭 수집"),
            ("Grafana", "메트릭 시각화 대시보드"),
            ("Sentry", "에러 추적 및 알림"),
        ],
    )

    # 12. 성능 목표
    doc.add_heading("12. 성능 목표", 1)
    make_table(
        doc,
        ["요청 유형", "p50 목표", "p95 목표"],
        [
            ("설정/상태 확인", "0.1~0.3s", "≤0.8s"),
            ("매치/랭킹 조회", "0.3~1s", "≤2s"),
            ("관리자 대시보드 조회", "0.3~1s", "≤2s"),
            ("사용량 조회", "0.1~0.5s", "≤1s"),
        ],
    )
    for n in [
        "동시 접속: 프로토타입 기준 최대 10명",
        "LLM 호출 RTT: RunPod ~150ms (콜드 스타트 ~15초), OpenAI ~500ms~2s",
        "SSE 연결 유지: 매치 진행 중 (~10~30분) 브라우저 연결 유지",
    ]:
        doc.add_paragraph(n, style="List Bullet")

    # 13. 배포 구성
    doc.add_heading("13. 배포 구성", 1)
    make_table(
        doc,
        ["항목", "값"],
        [
            ("인스턴스 타입", "t4g.small"),
            ("리전", "ap-northeast-2 (서울)"),
            ("배포 경로", "/opt/chatbot"),
            ("배포 방식", "Docker Compose (이미지 빌드 방식, 소스코드 COPY 베이킹)"),
            ("SSH 키", "~/Downloads/chatbot-key.pem"),
        ],
    )
    doc.add_heading("Docker Compose 서비스", 2)
    make_table(
        doc,
        ["서비스", "역할"],
        [
            ("backend", "FastAPI 애플리케이션"),
            ("frontend", "Next.js 애플리케이션"),
            ("postgres", "PostgreSQL 16"),
            ("redis", "Redis"),
        ],
    )
    doc.add_heading("배포 주의사항", 2)
    for n in [
        "소스코드 변경 시 반드시 이미지 재빌드 필요 (docker compose build)",
        "scp 파일 복사 후 restart만으로는 변경 사항 적용 불가",
        "uploads/ 디렉토리만 볼륨 마운트, 나머지는 이미지에 베이킹",
    ]:
        doc.add_paragraph(n, style="List Bullet")
    doc.add_heading("코드 변경 배포 명령", 2)
    code_block(
        doc,
        (
            "ssh -i ~/Downloads/chatbot-key.pem ubuntu@<EC2_IP> \\\n"
            '  "cd /opt/chatbot && git pull && \\\n'
            "   docker compose -f docker-compose.prod.yml build backend frontend && \\\n"
            '   docker compose -f docker-compose.prod.yml up -d backend frontend"'
        ),
    )

    # 변경 이력
    doc.add_heading("변경 이력", 1)
    make_table(
        doc,
        ["날짜", "버전", "변경 내용", "작성자"],
        [("2026-03-09", "v1.0", "최초 작성", "Claude")],
    )

    path = r"C:\Project_New\docs\시스템_아키텍처_설계서.docx"
    doc.save(path)
    print(f"저장 완료: {path}")


# ──────────────────────────────────────────────────────────────────
# 2. LLM 모델 아키텍처 설계서
# ──────────────────────────────────────────────────────────────────
def build_model_arch():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Title page
    t = doc.add_heading("LLM 모델 아키텍처 설계서", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("작성일: 2026-03-09    버전: v1.0    AI 에이전트 토론 플랫폼").italic = True
    doc.add_page_break()

    # TOC
    doc.add_heading("목차", 1)
    for item in [
        "1. 개요",
        "2. 모델 역할 분리",
        "3. LLM 라우팅 아키텍처",
        "4. 지원 모델 목록",
        "5. OptimizedOrchestrator 병렬 처리",
        "6. 모델 선정 벤치마크",
        "7. 토큰 사용량 추적",
        "8. 설정 값",
        "9. 플랫폼 크레딧 에이전트",
    ]:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # 1. 개요
    doc.add_heading("1. 개요", 1)
    doc.add_paragraph(
        "AI 에이전트 토론 플랫폼의 LLM 모델 아키텍처는 세 가지 역할로 분리된다."
    )
    make_table(
        doc,
        ["역할", "용도", "선택 방식"],
        [
            ("에이전트 발언 생성", "토론 참가 에이전트의 발언 생성", "에이전트 소유자가 직접 선택"),
            ("턴 검토 (Review)", "매 발언의 논리 오류·허위 주장·주제 이탈 탐지", "gpt-5-nano 고정"),
            ("최종 판정 (Judge)", "토론 종료 후 승자 결정", "gpt-4.1 고정"),
        ],
    )
    doc.add_paragraph(
        "모든 LLM 호출은 inference_client.py의 InferenceClient를 단일 진입점으로 사용한다. "
        "직접 openai.AsyncOpenAI() 등을 호출하는 것은 금지한다."
    )
    p = doc.add_paragraph()
    p.add_run("FigJam 다이어그램: ").bold = True
    p.add_run(
        "https://www.figma.com/online-whiteboard/create-diagram/b8595916-21c1-4b7b-8140-1746b11705fe"
    )

    # 2. 모델 역할 분리
    doc.add_heading("2. 모델 역할 분리", 1)
    doc.add_heading("역할별 모델 선정 근거", 2)
    make_table(
        doc,
        ["역할", "모델", "선정 근거"],
        [
            (
                "에이전트 발언",
                "에이전트별 선택",
                "사용자가 에이전트 성격·비용에 맞게 직접 선택",
            ),
            (
                "턴 검토 (Review)",
                "gpt-5-nano (고정)",
                "고속·저비용, 벌점 정확도 8.91점, 비용 $0.00017/1K tokens",
            ),
            (
                "최종 판정 (Judge)",
                "gpt-4.1 (고정)",
                "최고 판정 품질 8.94점, 비용 $0.0120/1K tokens",
            ),
        ],
    )
    doc.add_heading("Review 모델 역할 상세", 2)
    for item in [
        "검토 항목: 논리 오류, 허위 주장, 주제 이탈, 프롬프트 인젝션, 인신공격, 허수아비 논증, 순환논증, 성급한 일반화, 강조의 오류, 유전적 오류, 부적절한 호소, 미끄러운 경사, 분할/합성의 오류",
        "출력: 논리 점수(0~10) + 위반 항목별 벌점",
        "regex 기반 벌점 키: prompt_injection, ad_hominem 등 (접두사 없음)",
        "LLM 검토 기반 벌점 키: llm_prompt_injection, llm_ad_hominem, llm_straw_man, llm_circular_reasoning, llm_hasty_generalization, llm_accent, llm_genetic_fallacy, llm_appeal, llm_slippery_slope, llm_division, llm_composition, llm_off_topic, llm_false_claim",
        "실행 방식: 모든 발언에 항상 실행 (fast-path 스킵 없음)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Judge 모델 역할 상세", 2)
    for item in [
        "입력: 전체 턴 로그, 누적 벌점, 에이전트 프로필",
        "출력: 승자 ID, 판정 이유, 각 에이전트 최종 점수",
        "실행 시점: 턴 루프 완료 후 1회 호출",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 3. LLM 라우팅 아키텍처
    doc.add_heading("3. LLM 라우팅 아키텍처", 1)
    code_block(
        doc,
        (
            "서비스 → InferenceClient.generate()\n"
            "    → llm_models 테이블에서 provider/model_id 조회\n"
            "    → provider별 분기 (openai / anthropic / google / runpod / local)\n"
            "    → Langfuse 트레이스 기록\n"
            "    → token_usage_logs INSERT\n"
            "    → 응답 반환"
        ),
    )
    doc.add_heading("llm_models 테이블 주요 필드", 2)
    make_table(
        doc,
        ["필드", "설명"],
        [
            ("provider", "openai / anthropic / google / runpod / local"),
            ("model_id", "API 호출 시 사용하는 실제 모델 ID"),
            ("display_name", "UI에 표시되는 이름"),
            ("input_cost_per_1m", "입력 토큰 비용 (USD/1M tokens)"),
            ("output_cost_per_1m", "출력 토큰 비용 (USD/1M tokens)"),
            ("max_context_length", "최대 컨텍스트 길이 (tokens)"),
            ("is_active", "활성화 여부 (비활성화 시 사용 불가)"),
            ("tier", "모델 등급 (free / standard / premium)"),
        ],
    )

    # 4. 지원 모델 목록
    doc.add_heading("4. 지원 모델 목록", 1)
    doc.add_heading("OpenAI", 2)
    make_table(
        doc,
        ["모델", "용도"],
        [
            ("gpt-5.2", "최고 성능"),
            ("gpt-5.2-pro", "최고 성능 Pro"),
            ("gpt-5.1", "고성능"),
            ("gpt-5", "표준"),
            ("gpt-5-mini", "경량"),
            ("gpt-4.1", "Judge 기본값, 고판정 품질"),
            ("gpt-4.1-mini", "중간 성능/비용"),
            ("gpt-4.1-nano", "경량"),
            ("gpt-4o", "기존 표준"),
            ("gpt-4o-mini", "기존 경량"),
            ("o3", "추론 특화"),
            ("o3-pro", "추론 특화 Pro"),
            ("o4-mini", "추론 경량"),
        ],
    )
    doc.add_heading("Anthropic", 2)
    make_table(
        doc,
        ["모델", "용도"],
        [
            ("Claude Opus 4.6", "최고 성능"),
            ("Claude Sonnet 4.6", "균형 성능/비용"),
            ("Claude Haiku 4.5", "경량 고속"),
            ("Claude Sonnet 4.5", "이전 세대 표준"),
            ("Claude Opus 4.5", "이전 세대 최고"),
        ],
    )
    doc.add_heading("Google", 2)
    make_table(
        doc,
        ["모델", "용도"],
        [
            ("Gemini 3.1 Pro Preview", "최신 최고 성능"),
            ("Gemini 3 Flash Preview", "최신 경량"),
            ("Gemini 2.5 Pro", "고성능"),
            ("Gemini 2.5 Flash", "고속 경량"),
            ("Gemini 2.5 Flash-Lite", "초경량"),
        ],
    )
    doc.add_heading("RunPod Serverless (자체 호스팅)", 2)
    make_table(
        doc,
        ["모델", "파라미터", "용도"],
        [
            ("Llama 3.3 70B", "70B", "기본 추론 모델"),
            ("Llama 3.1 70B", "70B", "대체 추론 모델"),
            ("Mixtral 8x7B", "46.7B(활성)", "혼합 전문가"),
            ("Qwen 2.5 72B", "72B", "다국어 특화"),
        ],
    )
    for n in [
        "추론 엔진: SGLang",
        "콜드 스타트: ~15초 (Serverless 특성)",
        "EC2와 RTT: ~150ms",
    ]:
        doc.add_paragraph(n, style="List Bullet")
    doc.add_heading("Local (WebSocket)", 2)
    doc.add_paragraph(
        "사용자 PC에서 실행하는 Ollama 등 로컬 LLM을 연결한다. "
        "WebSocket으로 실시간 통신하며, API 키 불필요."
    )

    # 5. OptimizedOrchestrator
    doc.add_heading("5. OptimizedOrchestrator 병렬 처리", 1)
    doc.add_heading("처리 흐름", 2)
    doc.add_paragraph(
        "턴 루프에서 A 에이전트 발언 검토와 B 에이전트 발언 생성을 "
        "asyncio.gather()로 동시 실행한다."
    )
    code_block(
        doc,
        (
            "# asyncio.gather로 A 검토와 B 실행을 병렬화 — 턴 지연 37% 단축\n"
            "review_a, result_b = await asyncio.gather(\n"
            "    orchestrator.review_turn(turn_a),\n"
            "    engine.execute_turn(agent_b)\n"
            ")"
        ),
    )
    doc.add_heading("최적화 효과", 2)
    make_table(
        doc,
        ["항목", "순차 처리", "병렬 처리", "개선"],
        [
            ("턴당 소요 시간", "기준", "-37%", "37% 단축"),
            ("LLM 호출 비용", "기준", "-76%", "76% 절감"),
            ("LLM 호출 횟수", "기준", "-83%", "83% 감소"),
        ],
    )
    doc.add_heading("설정 플래그", 2)
    doc.add_paragraph(
        "config.py에서 debate_orchestrator_optimized = True로 활성화. "
        "False로 설정 시 즉시 순차 처리 방식으로 롤백."
    )

    # 6. 벤치마크
    doc.add_heading("6. 모델 선정 벤치마크", 1)
    doc.add_paragraph(
        "2026-02-26 수행한 GPT 전 모델 비교 벤치마크 결과를 바탕으로 "
        "Review/Judge 모델을 선정했다."
    )
    doc.add_heading("변경 전후 비교", 2)
    make_table(
        doc,
        ["구분", "항목", "기존", "최적화 후", "개선"],
        [
            ("Review", "모델", "gpt-4o-mini", "gpt-5-nano", "교체"),
            ("Review", "비용/1K tokens", "$0.0003", "$0.00017", "43% 절감"),
            ("Review", "품질 점수", "8.60점", "8.91점", "성능 향상"),
            ("Judge", "모델", "gpt-4o", "gpt-4.1", "교체"),
            ("Judge", "비용/1K tokens", "$0.0150", "$0.0120", "20% 절감"),
            ("Judge", "품질 점수", "8.50점", "8.94점", "성능 향상"),
            ("전체", "총 비용/매치", "$0.01739", "$0.01329", "23.6% 절감"),
        ],
    )
    doc.add_heading("Review 모델 후보 비교 (상위 3개)", 2)
    make_table(
        doc,
        ["순위", "모델", "품질 점수", "입력 비용/1M", "선정"],
        [
            ("1위", "gpt-5-mini", "9.12점", "$0.00040", "-"),
            ("2위 (선정)", "gpt-5-nano", "8.91점", "$0.00017", "선정"),
            ("3위", "gpt-4o-mini (기존)", "8.60점", "$0.00030", "교체"),
        ],
    )
    doc.add_heading("Judge 모델 후보 비교 (상위 5개)", 2)
    make_table(
        doc,
        ["순위", "모델", "품질 점수", "입력 비용/1M", "선정"],
        [
            ("1위", "gpt-5.2", "9.45점", "$0.01500", "-"),
            ("2위", "gpt-5", "9.21점", "$0.01300", "-"),
            ("3위", "o3", "9.18점", "$0.01000", "-"),
            ("4위", "gpt-5-mini", "9.07점", "$0.00080", "-"),
            ("5위 (선정)", "gpt-4.1", "8.94점", "$0.01200", "선정"),
            ("기존", "gpt-4o", "8.50점", "$0.01500", "교체"),
        ],
    )
    doc.add_paragraph(
        "관련 문서: docs/orchestrator_optimization.md    "
        "벤치마크 코드: backend/tests/benchmark/test_gpt_model_comparison.py"
    )

    # 7. 토큰 사용량 추적
    doc.add_heading("7. 토큰 사용량 추적", 1)
    doc.add_paragraph(
        "모든 LLM 호출 후 token_usage_logs 테이블에 즉시 기록한다. 기록 누락은 허용하지 않는다."
    )
    make_table(
        doc,
        ["필드", "설명"],
        [
            ("user_id", "호출을 발생시킨 사용자"),
            ("agent_id", "발언을 생성한 에이전트 (Review/Judge는 null)"),
            ("model_id", "사용한 LLM 모델"),
            ("role", "호출 역할 (agent_turn / review / judge)"),
            ("input_tokens", "입력 토큰 수"),
            ("output_tokens", "출력 토큰 수"),
            ("cost_usd", "비용 (USD)"),
            ("created_at", "기록 시각"),
        ],
    )
    doc.add_heading("조회 경로", 2)
    make_table(
        doc,
        ["역할", "API 경로"],
        [
            ("일반 사용자 (자신의 사용량)", "GET /api/usage/"),
            ("관리자 (전체 현황)", "GET /api/admin/usage/"),
            ("관리자 (모니터링 대시보드)", "GET /api/admin/monitoring/"),
        ],
    )
    doc.add_paragraph(
        "InferenceClient는 Langfuse에 비동기로 트레이스를 기록한다. "
        "Langfuse가 응답하지 않아도 LLM 호출 자체는 차단하지 않는다."
    )

    # 8. 설정 값
    doc.add_heading("8. 설정 값", 1)
    doc.add_paragraph("backend/app/core/config.py에서 관리하는 LLM 관련 설정:")
    make_table(
        doc,
        ["설정 키", "기본값", "설명"],
        [
            ("debate_review_model", '"gpt-5-nano"', "턴 검토에 사용할 모델"),
            ("debate_judge_model", '"gpt-4.1"', "최종 판정에 사용할 모델"),
            ("debate_orchestrator_optimized", "True", "병렬 처리 활성화 여부"),
            ("debate_turn_review_enabled", "True", "턴 검토 기능 활성화 여부"),
            ("debate_turn_review_timeout", "10", "검토 LLM 호출 타임아웃 (초)"),
        ],
    )
    doc.add_heading("롤백 방법", 2)
    doc.add_paragraph(
        "병렬 처리 오류 발생 시 환경 변수 DEBATE_ORCHESTRATOR_OPTIMIZED=false로 "
        "즉시 순차 처리 방식으로 복귀한다. 코드 변경 없이 적용 가능하다."
    )

    # 9. 플랫폼 크레딧 에이전트
    doc.add_heading("9. 플랫폼 크레딧 에이전트", 1)
    doc.add_paragraph(
        "에이전트 생성 시 use_platform_credits=True를 설정하면 사용자 개인 API 키 없이 "
        "플랫폼 공용 API 키를 사용해 LLM을 호출한다."
    )
    make_table(
        doc,
        ["항목", "일반 에이전트", "플랫폼 크레딧 에이전트"],
        [
            ("API 키", "사용자가 직접 입력 (BYOK)", "불필요"),
            ("비용 부담", "사용자 API 계정", "플랫폼 크레딧 차감"),
            ("매칭 큐 등록", "API 키 유효성 검사", "검사 생략"),
            ("LLM 호출 경로", "사용자 API 키 → 프로바이더", "플랫폼 env 키 → 프로바이더"),
        ],
    )
    doc.add_paragraph("API 키 저장 시 core/encryption.py의 Fernet 암호화를 적용한다.")

    # 변경 이력
    doc.add_heading("변경 이력", 1)
    make_table(
        doc,
        ["날짜", "버전", "변경 내용", "작성자"],
        [
            (
                "2026-03-09",
                "v1.0",
                "최초 작성 (벤치마크 결과, 병렬 처리, 라우팅 아키텍처 포함)",
                "Claude",
            )
        ],
    )

    path = r"C:\Project_New\docs\모델_아키텍처_설계서.docx"
    doc.save(path)
    print(f"저장 완료: {path}")


def build_combined():
    """시스템 아키텍처 + LLM 모델 아키텍처를 하나의 docx로 합본."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── 표지 ──
    t = doc.add_heading("AI 에이전트 토론 플랫폼\n설계 문서", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("시스템 아키텍처 설계서  ·  LLM 모델 아키텍처 설계서").bold = True
    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("작성일: 2026-03-09    버전: v1.0").italic = True
    doc.add_page_break()

    # ── 목차 ──
    doc.add_heading("목차", 1)
    toc = [
        ("1부  시스템 아키텍처", [
            "1. 서비스 개요",
            "2. 인프라 구성",
            "3. 전체 시스템 아키텍처",
            "4. 백엔드 구조",
            "5. 프론트엔드 구조",
            "6. 데이터베이스 구조",
            "7. 토론 엔진 흐름",
            "8. API 엔드포인트 명세",
            "9. 사용자 역할 및 접근 제어",
            "10. 실시간 스트리밍 구조",
            "11. 관측성 및 모니터링",
            "12. 성능 목표",
            "13. 배포 구성",
        ]),
        ("2부  LLM 모델 아키텍처", [
            "14. 개요",
            "15. 모델 역할 분리",
            "16. LLM 라우팅 아키텍처",
            "17. 지원 모델 목록",
            "18. OptimizedOrchestrator 병렬 처리",
            "19. 모델 선정 벤치마크",
            "20. 토큰 사용량 추적",
            "21. 설정 값",
            "22. 플랫폼 크레딧 에이전트",
        ]),
    ]
    for part_title, items in toc:
        p = doc.add_paragraph(part_title)
        p.runs[0].bold = True
        p.paragraph_format.space_before = Pt(8)
        for item in items:
            pi = doc.add_paragraph("    " + item)
            pi.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 1부: 시스템 아키텍처
    # ══════════════════════════════════════════════════════
    part1 = doc.add_heading("1부  시스템 아키텍처", 1)
    part1.runs[0].font.size = Pt(18)
    doc.add_paragraph()

    # 1. 서비스 개요
    doc.add_heading("1. 서비스 개요", 2)
    doc.add_paragraph(
        "AI 에이전트 토론 플랫폼은 AI 에이전트끼리 실시간으로 토론을 벌이고, "
        "사용자가 관전·예측투표·시즌 랭킹을 즐기는 플랫폼이다."
    )
    make_table(doc, ["항목", "내용"], [
        ("서비스 단계", "프로토타입 (동시 접속 10명 이하)"),
        ("월 예상 비용", "~$130 (EC2 ~$15 + RunPod ~$114) + LLM API 비용 (사용량 비례)"),
        ("주요 특징", "AI 에이전트 토론, 실시간 관전, ELO 랭킹, 예측투표, 토너먼트, 턴 검토 시스템"),
    ])
    doc.add_heading("핵심 기능", 3)
    for f in [
        "AI 에이전트 토론: 사용자가 에이전트(성격·모델·프롬프트)를 직접 생성하고 토론에 참가",
        "실시간 관전: Redis Pub/Sub + SSE로 토론 진행 상황 실시간 브로드캐스트",
        "턴 검토 시스템: gpt-5-nano가 매 발언을 검토 (논리 오류·허위 주장·주제 이탈 탐지, 벌점 부여)",
        "ELO 랭킹·시즌: ELO 기반 시즌 랭킹, 승급전(3판 2선승)/강등전(1판) 자동 생성",
        "예측투표: 매치 시작 전 사용자 승자 예측, 완료 후 결과 공개",
        "토너먼트: 대진표 자동 생성, 단계별 진행",
        "LLM 모델 전환: 에이전트별 LLM 모델 선택 가능 (OpenAI/Anthropic/Google/RunPod/Local)",
        "토큰 사용량 추적: 사용자별 LLM 토큰 사용량 실시간 추적 및 비용 산출",
        "관리자 대시보드: 매치 관리, 시즌/토너먼트 관리, 모니터링, 사용량/과금 현황",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    # 2. 인프라 구성
    doc.add_heading("2. 인프라 구성", 2)
    make_table(doc, ["역할", "기술"], [
        ("Frontend", "Next.js 15 + React 19 + Zustand, Tailwind CSS"),
        ("Backend", "Python 3.12 + FastAPI, SQLAlchemy 2.0 async"),
        ("Database", "PostgreSQL 16 (Docker) — 18개 테이블"),
        ("Cache/Pub-Sub", "Redis (Docker)"),
        ("LLM Inference", "RunPod Serverless + SGLang (기본) + OpenAI/Anthropic/Google API"),
        ("Streaming", "SSE (Server-Sent Events)"),
        ("Observability", "Langfuse + Prometheus + Grafana + Sentry"),
        ("Infra", "AWS EC2 t4g.small (서울 ap-northeast-2) + RunPod Serverless (미국)"),
        ("Container", "Docker Compose"),
    ])
    for d in [
        "EC2 t4g.small (서울): FastAPI 백엔드, PostgreSQL, Redis를 Docker Compose로 함께 운용",
        "RunPod Serverless (미국): Llama 3.3 70B SGLang 추론 서버, 콜드 스타트 ~15초",
        "외부 LLM API: OpenAI, Anthropic, Google — RTT ~150ms (에이전트 발언 생성 시)",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    # 3. 전체 시스템 아키텍처
    doc.add_heading("3. 전체 시스템 아키텍처", 2)
    code_block(doc, (
        "┌─ 사용자 브라우저 ─────────────────────────────────┐\n"
        "│  Next.js 15 Frontend  (Zustand + SSE Client)    │\n"
        "└──────────────────┬────────────────────────────────┘\n"
        "                   │ HTTPS / SSE / WebSocket\n"
        "┌──────────────────▼────────────────────────────────┐\n"
        "│  EC2 서울 t4g.small                               │\n"
        "│  ┌─ FastAPI ───────────────────────────────────┐  │\n"
        "│  │  /api/auth   /api/agents   /api/matches      │  │\n"
        "│  │  /api/topics /api/tournaments /api/ws         │  │\n"
        "│  │  /api/models /api/usage   /api/admin          │  │\n"
        "│  └─────────────────────────────────────────────┘  │\n"
        "│  PostgreSQL 16 (18 tables)  Redis Pub/Sub + Cache │\n"
        "└──────────────────┬────────────────────────────────┘\n"
        "                   │ LLM HTTP (~150ms RTT)\n"
        "┌──────────────────▼────────────────────────────────┐\n"
        "│  LLM 프로바이더                                   │\n"
        "│  OpenAI API  /  Anthropic API  /  Google API      │\n"
        "│  RunPod Serverless (Llama 3.3 70B / SGLang)       │\n"
        "└───────────────────────────────────────────────────┘"
    ))
    p = doc.add_paragraph()
    p.add_run("FigJam 다이어그램: ").bold = True
    p.add_run("https://www.figma.com/online-whiteboard/create-diagram/ccd7315e-f8c2-44a1-813f-1b6423ceca7a")

    # 4. 백엔드 구조
    doc.add_heading("4. 백엔드 구조", 2)
    doc.add_heading("디렉토리 구조", 3)
    code_block(doc, (
        "backend/app/\n"
        "├── main.py                  # FastAPI 앱, 라우터 등록\n"
        "├── api/                     # 라우터 레이어 (입력 검증 + HTTP 응답만)\n"
        "│   ├── auth.py\n"
        "│   ├── debate_agents.py\n"
        "│   ├── debate_matches.py\n"
        "│   ├── debate_topics.py\n"
        "│   ├── debate_tournaments.py\n"
        "│   ├── debate_ws.py         # WebSocket (로컬 에이전트 연결)\n"
        "│   └── admin/\n"
        "│       ├── debate/          # 매치·시즌·토너먼트·에이전트 관리\n"
        "│       └── system/          # 사용자·LLM 모델·모니터링·사용량 관리\n"
        "├── core/                    # 인프라 설정\n"
        "├── models/                  # SQLAlchemy ORM 모델 (18개)\n"
        "├── schemas/                 # Pydantic v2 입출력 스키마\n"
        "└── services/                # 비즈니스 로직"
    ))
    doc.add_heading("주요 서비스 목록", 3)
    make_table(doc, ["파일", "역할"], [
        ("debate_agent_service.py", "에이전트 CRUD, 랭킹, 갤러리, 클론, H2H, 버전 관리"),
        ("debate_match_service.py", "매치 조회, 하이라이트, 요약 리포트 생성"),
        ("debate_matching_service.py", "큐 등록/취소, 자동 매칭(DebateAutoMatcher), ready_up"),
        ("debate_engine.py", "토론 실행 루프 (턴 실행 → 검토 → 판정 → 결과 저장)"),
        ("debate_orchestrator.py", "LLM 검토 + 최적화 병렬 실행(OptimizedDebateOrchestrator)"),
        ("debate_broadcast.py", "SSE 이벤트 발행/구독, 관전자 수 관리"),
        ("debate_ws_manager.py", "WebSocket 연결 관리 (로컬 에이전트 인증·메시지 라우팅)"),
        ("inference_client.py", "LLM 호출 단일 진입점 (Langfuse 추적, 토큰 로깅, provider 분기)"),
        ("debate_season_service.py", "시즌 생성/종료, 시즌 ELO 집계, 보상 지급"),
        ("debate_promotion_service.py", "승급전/강등전 시리즈 생성·진행·완료 처리"),
        ("debate_tournament_service.py", "토너먼트 대진표 생성·진행"),
    ])
    doc.add_heading("레이어 책임 원칙", 3)
    make_table(doc, ["레이어", "책임", "금지 사항"], [
        ("api/ (라우터)", "HTTP 요청 수신, 입력 검증, 응답 포맷", "DB 직접 쿼리, 비즈니스 로직"),
        ("services/", "비즈니스 로직, DB 조작", "HTTP 관련 코드"),
        ("models/", "ORM 정의, 테이블 구조", "비즈니스 로직"),
        ("schemas/", "입출력 데이터 검증", "DB 접근"),
        ("core/", "인프라 설정, 공통 의존성", "도메인 로직"),
    ])

    # 5. 프론트엔드 구조
    doc.add_heading("5. 프론트엔드 구조", 2)
    code_block(doc, (
        "frontend/src/\n"
        "├── app/\n"
        "│   ├── (user)/              # 사용자 라우트 그룹\n"
        "│   │   ├── debate/          # 토론 목록, 매치 관전, 갤러리, 랭킹\n"
        "│   │   ├── agents/          # 에이전트 생성/편집\n"
        "│   │   ├── seasons/         # 시즌 랭킹\n"
        "│   │   ├── tournaments/     # 토너먼트 대진표\n"
        "│   │   └── mypage/          # 내 정보, 사용량\n"
        "│   ├── admin/               # 관리자 라우트 그룹\n"
        "│   └── api/[...path]/       # Next.js → FastAPI SSE 프록시\n"
        "├── components/\n"
        "│   ├── debate/              # DebateViewer, TurnBubble, PromotionBadge 등\n"
        "│   └── layout/              # 공통 레이아웃\n"
        "├── stores/                  # Zustand 상태 관리\n"
        "└── lib/\n"
        "    ├── api.ts               # API 호출 유틸리티\n"
        "    ├── auth.ts              # 인증 처리\n"
        "    └── agentWebSocket.ts    # 로컬 에이전트 WebSocket 클라이언트"
    ))
    make_table(doc, ["상태 (debateStore.ts)", "설명"], [
        ("currentMatch", "현재 관전 중인 매치 정보"),
        ("turns", "누적된 턴 발언 목록"),
        ("turnReviews", "턴별 LLM 검토 결과 (벌점, 논리 점수)"),
        ("prediction", "사용자의 예측투표 상태"),
        ("viewerCount", "현재 관전자 수"),
        ("replayMode / replayIndex", "리플레이 모드 상태"),
    ])

    # 6. 데이터베이스 구조
    doc.add_heading("6. 데이터베이스 구조", 2)
    make_table(doc, ["모델", "테이블", "설명"], [
        ("User", "users", "사용자 계정, 역할(user/admin/superadmin), 크레딧 잔액"),
        ("LLMModel", "llm_models", "등록된 LLM 모델 (provider, 비용, 활성화 여부)"),
        ("TokenUsageLog", "token_usage_logs", "LLM 호출 토큰·비용 기록"),
        ("DebateAgent", "debate_agents", "에이전트 (소유자, provider, ELO, 공개 여부, 승급전 상태)"),
        ("DebateAgentVersion", "debate_agent_versions", "에이전트 버전 이력 (system_prompt 스냅샷)"),
        ("DebateAgentSeasonStats", "debate_agent_season_stats", "시즌별 ELO·전적 분리 집계"),
        ("DebateAgentTemplate", "debate_agent_templates", "관리자 제공 에이전트 템플릿"),
        ("DebateTopic", "debate_topics", "토론 주제 (등록자, 승인 상태)"),
        ("DebateMatch", "debate_matches", "매치 (참가자, 형식, 상태, 결과, 시즌/시리즈 연결)"),
        ("DebateMatchParticipant", "debate_match_participants", "멀티에이전트 매치 참가자 목록"),
        ("DebateMatchPrediction", "debate_match_predictions", "사용자 예측투표"),
        ("DebateMatchQueue", "debate_match_queues", "매칭 대기 큐"),
        ("DebateTurnLog", "debate_turn_logs", "턴별 발언·검토 결과·점수 기록"),
        ("DebatePromotionSeries", "debate_promotion_series", "승급전/강등전 시리즈 상태"),
        ("DebateSeason", "debate_seasons", "시즌 기간·상태"),
        ("DebateSeasonResult", "debate_season_results", "시즌 종료 시 최종 순위 스냅샷"),
        ("DebateTournament", "debate_tournaments", "토너먼트 대진표·상태"),
        ("DebateTournamentEntry", "debate_tournament_entries", "토너먼트 참가 에이전트 목록"),
    ])

    # 7. 토론 엔진 흐름
    doc.add_heading("7. 토론 엔진 흐름", 2)
    code_block(doc, (
        "큐 등록 → DebateAutoMatcher 감지 → ready_up() → DebateMatch 생성\n"
        "    → debate_engine.run_match()\n"
        "        ├─ 턴 루프 (N 라운드)\n"
        "        │   ├─ 에이전트 발언 생성 (LLM 호출 or WebSocket)\n"
        "        │   └─ OptimizedDebateOrchestrator.review_turn()\n"
        "        │       └─ asyncio.gather(A 검토, B 실행) 병렬 실행\n"
        "        └─ judge() → 최종 판정 → ELO 갱신 → 승급전 체크\n"
        "    → SSE 이벤트 발행 (debate_broadcast → Redis Pub/Sub)"
    ))
    p = doc.add_paragraph()
    p.add_run("FigJam 시퀀스 다이어그램: ").bold = True
    p.add_run("https://www.figma.com/online-whiteboard/create-diagram/fdf2f969-2ca9-4695-8ca9-6ed628c857d0")
    make_table(doc, ["종류", "조건", "규칙"], [
        ("승급전", "ELO 일정 임계값 초과", "3판 2선승 (required_wins=2)"),
        ("강등전", "ELO 일정 임계값 미만", "1판 필승 (required_wins=1)"),
    ])

    # 8. API 엔드포인트 명세
    doc.add_heading("8. API 엔드포인트 명세", 2)
    make_table(doc, ["경로", "파일", "설명"], [
        ("GET /health", "health.py", "서버 상태 확인"),
        ("/api/auth/*", "auth.py", "회원가입, 로그인, 토큰 갱신"),
        ("/api/agents/*", "debate_agents.py", "에이전트 CRUD, 랭킹, 갤러리, H2H"),
        ("/api/topics/*", "debate_topics.py", "토픽 등록/조회/매칭 큐"),
        ("/api/matches/*", "debate_matches.py", "매치 조회, SSE 스트리밍, 예측투표, 요약"),
        ("/api/tournaments/*", "debate_tournaments.py", "토너먼트 CRUD, 대진표"),
        ("/api/models/*", "models.py", "LLM 모델 목록 조회, 선호 모델 설정"),
        ("/api/usage/*", "usage.py", "내 토큰 사용량 조회"),
        ("/api/ws/debate/*", "debate_ws.py", "WebSocket (로컬 에이전트 전용)"),
        ("/api/admin/users/*", "admin/system/users.py", "사용자 조회/역할 변경"),
        ("/api/admin/models/*", "admin/system/llm_models.py", "LLM 모델 등록/수정/활성화"),
        ("/api/admin/usage/*", "admin/system/usage.py", "전체 사용량 현황"),
        ("/api/admin/monitoring/*", "admin/system/monitoring.py", "토큰/비용 모니터링"),
        ("/api/admin/debate/*", "admin/debate/", "매치 강제실행, 시즌/토너먼트 관리"),
    ])
    doc.add_heading("SSE 스트리밍 이벤트", 3)
    make_table(doc, ["이벤트 타입", "설명"], [
        ("match_start", "토론 시작 알림"),
        ("turn", "에이전트 발언"),
        ("turn_review", "턴 검토 결과 (벌점, 논리 점수)"),
        ("match_end", "토론 완료, 최종 결과"),
        ("series_update", "승급전/강등전 시리즈 상태 변경"),
        ("viewer_count", "현재 관전자 수 갱신"),
    ])

    # 9. RBAC
    doc.add_heading("9. 사용자 역할 및 접근 제어 (RBAC)", 2)
    make_table(doc, ["역할", "접근 범위", "주요 기능"], [
        ("user", "토론 관전, 에이전트 생성/편집, 예측투표, 랭킹 조회, 사용량 조회",
         "에이전트 커스터마이징, 큐 등록, 토너먼트 참가, LLM 모델 선택"),
        ("admin", "관리자 대시보드 + 사용자 화면 전체 (읽기 위주)",
         "매치 관리, 시즌/토너먼트 관리, 모니터링, 에이전트 모더레이션"),
        ("superadmin", "admin 전체 + 파괴적 작업",
         "사용자 삭제/역할 변경, LLM 모델 등록/수정, 시스템 설정, 쿼터 관리"),
    ])

    # 10. 실시간 스트리밍
    doc.add_heading("10. 실시간 스트리밍 구조", 2)
    code_block(doc, (
        'DebateEngine\n'
        '    → Redis PUBLISH "debate:{match_id}"\n'
        '        → FastAPI SSE 구독자\n'
        '            → HTTP/2 SSE 스트림\n'
        '                → 브라우저 EventSource'
    ))

    # 11. 관측성
    doc.add_heading("11. 관측성 및 모니터링", 2)
    make_table(doc, ["도구", "역할"], [
        ("Langfuse", "LLM 호출 추적 (모델, 토큰, 응답 시간, 비용)"),
        ("Prometheus", "애플리케이션 메트릭 수집"),
        ("Grafana", "메트릭 시각화 대시보드"),
        ("Sentry", "에러 추적 및 알림"),
    ])

    # 12. 성능 목표
    doc.add_heading("12. 성능 목표", 2)
    make_table(doc, ["요청 유형", "p50 목표", "p95 목표"], [
        ("설정/상태 확인", "0.1~0.3s", "≤0.8s"),
        ("매치/랭킹 조회", "0.3~1s", "≤2s"),
        ("관리자 대시보드 조회", "0.3~1s", "≤2s"),
        ("사용량 조회", "0.1~0.5s", "≤1s"),
    ])

    # 13. 배포 구성
    doc.add_heading("13. 배포 구성", 2)
    make_table(doc, ["항목", "값"], [
        ("인스턴스 타입", "t4g.small"),
        ("리전", "ap-northeast-2 (서울)"),
        ("배포 경로", "/opt/chatbot"),
        ("배포 방식", "Docker Compose (이미지 빌드 방식, 소스코드 COPY 베이킹)"),
        ("SSH 키", "~/Downloads/chatbot-key.pem"),
    ])
    code_block(doc, (
        "ssh -i ~/Downloads/chatbot-key.pem ubuntu@<EC2_IP> \\\n"
        '  "cd /opt/chatbot && git pull && \\\n'
        "   docker compose -f docker-compose.prod.yml build backend frontend && \\\n"
        '   docker compose -f docker-compose.prod.yml up -d backend frontend"'
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 2부: LLM 모델 아키텍처
    # ══════════════════════════════════════════════════════
    part2 = doc.add_heading("2부  LLM 모델 아키텍처", 1)
    part2.runs[0].font.size = Pt(18)
    doc.add_paragraph()

    # 14. 개요
    doc.add_heading("14. 개요", 2)
    doc.add_paragraph("AI 에이전트 토론 플랫폼의 LLM 모델 아키텍처는 세 가지 역할로 분리된다.")
    make_table(doc, ["역할", "용도", "선택 방식"], [
        ("에이전트 발언 생성", "토론 참가 에이전트의 발언 생성", "에이전트 소유자가 직접 선택"),
        ("턴 검토 (Review)", "매 발언의 논리 오류·허위 주장·주제 이탈 탐지", "gpt-5-nano 고정"),
        ("최종 판정 (Judge)", "토론 종료 후 승자 결정", "gpt-4.1 고정"),
    ])
    doc.add_paragraph(
        "모든 LLM 호출은 inference_client.py의 InferenceClient를 단일 진입점으로 사용한다. "
        "직접 openai.AsyncOpenAI() 등을 호출하는 것은 금지한다."
    )
    p = doc.add_paragraph()
    p.add_run("FigJam 다이어그램: ").bold = True
    p.add_run("https://www.figma.com/online-whiteboard/create-diagram/b8595916-21c1-4b7b-8140-1746b11705fe")

    # 15. 모델 역할 분리
    doc.add_heading("15. 모델 역할 분리", 2)
    make_table(doc, ["역할", "모델", "선정 근거"], [
        ("에이전트 발언", "에이전트별 선택", "사용자가 에이전트 성격·비용에 맞게 직접 선택"),
        ("턴 검토 (Review)", "gpt-5-nano (고정)", "고속·저비용, 벌점 정확도 8.91점, 비용 $0.00017/1K tokens"),
        ("최종 판정 (Judge)", "gpt-4.1 (고정)", "최고 판정 품질 8.94점, 비용 $0.0120/1K tokens"),
    ])
    for item in [
        "검토 항목: 논리 오류, 허위 주장, 주제 이탈, 프롬프트 인젝션, 인신공격, 허수아비 논증, 순환논증, 성급한 일반화, 강조의 오류, 유전적 오류, 부적절한 호소, 미끄러운 경사, 분할/합성의 오류",
        "출력: 논리 점수(0~10) + 위반 항목별 벌점",
        "LLM 검토 벌점 키: llm_prompt_injection, llm_ad_hominem, llm_straw_man, llm_circular_reasoning, llm_hasty_generalization, llm_accent, llm_genetic_fallacy, llm_appeal, llm_slippery_slope, llm_division, llm_composition, llm_off_topic, llm_false_claim",
        "실행 방식: 모든 발언에 항상 실행 (fast-path 스킵 없음)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 16. LLM 라우팅 아키텍처
    doc.add_heading("16. LLM 라우팅 아키텍처", 2)
    code_block(doc, (
        "서비스 → InferenceClient.generate()\n"
        "    → llm_models 테이블에서 provider/model_id 조회\n"
        "    → provider별 분기 (openai / anthropic / google / runpod / local)\n"
        "    → Langfuse 트레이스 기록\n"
        "    → token_usage_logs INSERT\n"
        "    → 응답 반환"
    ))
    make_table(doc, ["필드", "설명"], [
        ("provider", "openai / anthropic / google / runpod / local"),
        ("model_id", "API 호출 시 사용하는 실제 모델 ID"),
        ("input_cost_per_1m", "입력 토큰 비용 (USD/1M tokens)"),
        ("output_cost_per_1m", "출력 토큰 비용 (USD/1M tokens)"),
        ("is_active", "활성화 여부 (비활성화 시 사용 불가)"),
        ("tier", "모델 등급 (free / standard / premium)"),
    ])

    # 17. 지원 모델 목록
    doc.add_heading("17. 지원 모델 목록", 2)
    doc.add_heading("OpenAI", 3)
    make_table(doc, ["모델", "용도"], [
        ("gpt-5.2 / gpt-5.2-pro", "최고 성능"),
        ("gpt-5.1 / gpt-5", "고성능 / 표준"),
        ("gpt-5-mini", "경량"),
        ("gpt-4.1", "Judge 기본값, 고판정 품질"),
        ("gpt-4.1-mini / gpt-4.1-nano", "중간/경량"),
        ("gpt-4o / gpt-4o-mini", "기존 세대"),
        ("o3 / o3-pro / o4-mini", "추론 특화"),
    ])
    doc.add_heading("Anthropic", 3)
    make_table(doc, ["모델", "용도"], [
        ("Claude Opus 4.6", "최고 성능"),
        ("Claude Sonnet 4.6", "균형 성능/비용"),
        ("Claude Haiku 4.5", "경량 고속"),
    ])
    doc.add_heading("Google", 3)
    make_table(doc, ["모델", "용도"], [
        ("Gemini 3.1 Pro Preview", "최신 최고 성능"),
        ("Gemini 2.5 Pro / Flash", "고성능 / 고속 경량"),
        ("Gemini 2.5 Flash-Lite", "초경량"),
    ])
    doc.add_heading("RunPod Serverless / Local", 3)
    make_table(doc, ["모델", "파라미터", "용도"], [
        ("Llama 3.3 70B", "70B", "기본 추론 모델 (SGLang, EC2 RTT ~150ms)"),
        ("Qwen 2.5 72B", "72B", "다국어 특화"),
        ("Local (WebSocket)", "-", "Ollama 등 온프레미스 LLM, API 키 불필요"),
    ])

    # 18. OptimizedOrchestrator
    doc.add_heading("18. OptimizedOrchestrator 병렬 처리", 2)
    doc.add_paragraph(
        "턴 루프에서 A 에이전트 발언 검토와 B 에이전트 발언 생성을 "
        "asyncio.gather()로 동시 실행한다."
    )
    code_block(doc, (
        "# asyncio.gather로 A 검토와 B 실행을 병렬화 — 턴 지연 37% 단축\n"
        "review_a, result_b = await asyncio.gather(\n"
        "    orchestrator.review_turn(turn_a),\n"
        "    engine.execute_turn(agent_b)\n"
        ")"
    ))
    make_table(doc, ["항목", "순차 처리", "병렬 처리", "개선"], [
        ("턴당 소요 시간", "기준", "-37%", "37% 단축"),
        ("LLM 호출 비용", "기준", "-76%", "76% 절감"),
        ("LLM 호출 횟수", "기준", "-83%", "83% 감소"),
    ])

    # 19. 벤치마크
    doc.add_heading("19. 모델 선정 벤치마크", 2)
    doc.add_paragraph("2026-02-26 수행한 GPT 전 모델 비교 벤치마크 결과를 바탕으로 Review/Judge 모델을 선정했다.")
    make_table(doc, ["구분", "항목", "기존", "최적화 후", "개선"], [
        ("Review", "모델", "gpt-4o-mini", "gpt-5-nano", "교체"),
        ("Review", "비용/1K tokens", "$0.0003", "$0.00017", "43% 절감"),
        ("Review", "품질 점수", "8.60점", "8.91점", "성능 향상"),
        ("Judge", "모델", "gpt-4o", "gpt-4.1", "교체"),
        ("Judge", "비용/1K tokens", "$0.0150", "$0.0120", "20% 절감"),
        ("Judge", "품질 점수", "8.50점", "8.94점", "성능 향상"),
        ("전체", "총 비용/매치", "$0.01739", "$0.01329", "23.6% 절감"),
    ])

    # 20. 토큰 사용량 추적
    doc.add_heading("20. 토큰 사용량 추적", 2)
    make_table(doc, ["필드", "설명"], [
        ("user_id", "호출을 발생시킨 사용자"),
        ("agent_id", "발언을 생성한 에이전트 (Review/Judge는 null)"),
        ("model_id", "사용한 LLM 모델"),
        ("role", "호출 역할 (agent_turn / review / judge)"),
        ("input_tokens / output_tokens", "토큰 수"),
        ("cost_usd", "비용 (USD)"),
        ("created_at", "기록 시각"),
    ])
    make_table(doc, ["역할", "API 경로"], [
        ("일반 사용자 (자신의 사용량)", "GET /api/usage/"),
        ("관리자 (전체 현황)", "GET /api/admin/usage/"),
        ("관리자 (모니터링)", "GET /api/admin/monitoring/"),
    ])

    # 21. 설정 값
    doc.add_heading("21. 설정 값 (config.py)", 2)
    make_table(doc, ["설정 키", "기본값", "설명"], [
        ("debate_review_model", '"gpt-5-nano"', "턴 검토에 사용할 모델"),
        ("debate_judge_model", '"gpt-4.1"', "최종 판정에 사용할 모델"),
        ("debate_orchestrator_optimized", "True", "병렬 처리 활성화 여부"),
        ("debate_turn_review_enabled", "True", "턴 검토 기능 활성화 여부"),
        ("debate_turn_review_timeout", "10", "검토 LLM 호출 타임아웃 (초)"),
    ])

    # 22. 플랫폼 크레딧 에이전트
    doc.add_heading("22. 플랫폼 크레딧 에이전트", 2)
    make_table(doc, ["항목", "일반 에이전트", "플랫폼 크레딧 에이전트"], [
        ("API 키", "사용자가 직접 입력 (BYOK)", "불필요"),
        ("비용 부담", "사용자 API 계정", "플랫폼 크레딧 차감"),
        ("매칭 큐 등록", "API 키 유효성 검사", "검사 생략"),
        ("LLM 호출 경로", "사용자 API 키 → 프로바이더", "플랫폼 env 키 → 프로바이더"),
    ])

    # 변경 이력
    doc.add_heading("변경 이력", 1)
    make_table(doc, ["날짜", "버전", "변경 내용", "작성자"], [
        ("2026-03-09", "v1.0", "시스템 아키텍처 + LLM 모델 아키텍처 합본 최초 작성", "Claude"),
    ])

    path = r"C:\Project_New\docs\AI_토론_플랫폼_설계서.docx"
    doc.save(path)
    print(f"저장 완료: {path}")


if __name__ == "__main__":
    build_combined()
    print("합본 문서 생성 완료")
