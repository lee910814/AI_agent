"""AI 토론 플랫폼 아키텍처 문서 — 다이어그램 이미지 포함"""
import io, textwrap
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 폰트 ──
FP      = fm.FontProperties(fname=r"C:\Windows\Fonts\malgun.ttf")
FP_BOLD = fm.FontProperties(fname=r"C:\Windows\Fonts\malgunbd.ttf")
matplotlib.rcParams["axes.unicode_minus"] = False

# ── 색상 ──
C_BLUE   = "#1F497D"; C_LBLUE  = "#4472C4"; C_TEAL   = "#17A589"
C_ORANGE = "#E67E22"; C_DARK_O = "#D35400"; C_GRAY   = "#5D6D7E"
C_LGRAY  = "#EAF0FB"; C_WHITE  = "#FFFFFF"; C_RED    = "#C0392B"
C_GREEN  = "#1E8449"; C_PURPLE = "#7D3C98"


def T(ax, x, y, s, bold=False, fs=9, color=C_WHITE,
      ha="center", va="center", zorder=4, **kw):
    """한글 폰트 보장 텍스트 헬퍼."""
    ax.text(x, y, s, ha=ha, va=va, fontsize=fs, color=color,
            fontproperties=FP_BOLD if bold else FP, zorder=zorder, **kw)


def buf(fig):
    b = io.BytesIO()
    fig.savefig(b, format="png", dpi=150, bbox_inches="tight")
    b.seek(0); plt.close(fig); return b


def bbox(ax, x, y, w, h, fc, ec=None, radius=0.2, zorder=2, lw=1.5):
    ax.add_patch(FancyBboxPatch(
        (x, y), w, h,
        boxstyle=f"round,pad=0.04,rounding_size={radius}",
        fc=fc, ec=ec or fc, linewidth=lw, zorder=zorder))


def arrow(ax, x1, y1, x2, y2, color=C_GRAY, lw=1.6):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color, lw=lw), zorder=6)


# ══════════════════════════════════════════════════════
# 다이어그램 1 — 전체 시스템 아키텍처
# ══════════════════════════════════════════════════════
def draw_system_arch():
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12); ax.set_ylim(0, 8); ax.axis("off")
    fig.patch.set_facecolor("#F0F3F7")

    # ─ 브라우저 그룹 ─
    bbox(ax, 0.2, 6.5, 3.5, 1.2, "#D5EEF8", ec=C_TEAL, lw=2, zorder=1, radius=0.25)
    T(ax, 0.45, 7.52, "사용자 브라우저", bold=True, fs=8, color=C_TEAL, ha="left", va="top")
    bbox(ax, 0.4, 6.65, 3.1, 0.8, C_TEAL, zorder=3)
    T(ax, 1.95, 7.1, "Next.js 15 Frontend", bold=True, fs=9)
    T(ax, 1.95, 6.85, "Zustand  ·  SSE Client", bold=False, fs=7.5)

    # ─ EC2 그룹 ─
    bbox(ax, 0.2, 2.0, 7.6, 4.2, "#E8EEF7", ec=C_BLUE, lw=2, zorder=1, radius=0.25)
    T(ax, 0.45, 6.0, "EC2 서울  (t4g.small)", bold=True, fs=8, color=C_BLUE, ha="left", va="top")

    # FastAPI 라우터 (2줄로 분리)
    bbox(ax, 0.4, 5.15, 7.2, 0.65, C_LBLUE, zorder=3)
    T(ax, 4.0, 5.6, "/api/auth  /api/agents  /api/matches  /api/topics", bold=True, fs=8)
    T(ax, 4.0, 5.33, "/api/tournaments  /api/models  /api/usage  /api/admin", bold=True, fs=8)

    # 서비스 박스들
    bbox(ax, 0.4, 4.1, 3.3, 0.85, C_BLUE, zorder=3)
    T(ax, 2.05, 4.62, "FastAPI Backend", bold=True, fs=9)
    T(ax, 2.05, 4.35, "비즈니스 로직 · SSE · WS", bold=False, fs=7.5)

    bbox(ax, 3.9, 4.1, 1.9, 0.85, C_GRAY, zorder=3)
    T(ax, 4.85, 4.62, "PostgreSQL 16", bold=True, fs=9)
    T(ax, 4.85, 4.35, "18개 테이블", bold=False, fs=7.5)

    bbox(ax, 5.95, 4.1, 1.65, 0.85, C_GRAY, zorder=3)
    T(ax, 6.77, 4.62, "Redis", bold=True, fs=9)
    T(ax, 6.77, 4.35, "Pub/Sub · Cache", bold=False, fs=7.5)

    # 관측성
    bbox(ax, 0.4, 2.2, 7.2, 1.65, "#6C3483", zorder=3)
    T(ax, 4.0, 3.27, "Observability", bold=True, fs=10)
    T(ax, 4.0, 2.95, "Langfuse", bold=False, fs=8.5)
    T(ax, 4.0, 2.65, "Prometheus  ·  Grafana  ·  Sentry", bold=False, fs=8.5)

    # ─ LLM 그룹 ─
    bbox(ax, 8.8, 2.0, 3.0, 4.2, "#FEF9E7", ec=C_ORANGE, lw=2, zorder=1, radius=0.25)
    T(ax, 9.05, 6.0, "LLM 프로바이더", bold=True, fs=8, color=C_ORANGE, ha="left", va="top")
    for i, (name, sub, fc) in enumerate([
        ("OpenAI API",  "gpt-5-nano · gpt-4.1 등", C_ORANGE),
        ("Anthropic",   "Claude Opus/Sonnet 4.6",  C_ORANGE),
        ("Google API",  "Gemini 3.1 Pro 등",        C_ORANGE),
        ("RunPod",      "Llama 3.3 70B · SGLang",  C_DARK_O),
    ]):
        py = 5.0 - i * 0.9
        bbox(ax, 9.0, py, 2.6, 0.7, fc, zorder=3)
        T(ax, 10.3, py+0.44, name, bold=True, fs=9)
        T(ax, 10.3, py+0.18, sub, bold=False, fs=7.5)

    # ─ 화살표 ─
    arrow(ax, 1.95, 6.65, 1.95, 5.8, C_TEAL)
    T(ax, 2.45, 6.2, "HTTPS/SSE/WS", bold=False, fs=7.5, color=C_TEAL)

    arrow(ax, 3.7, 4.52, 3.9, 4.52)
    arrow(ax, 5.8, 4.52, 5.95, 4.52)

    for i, py in enumerate([5.35, 4.45, 3.55, 2.65]):
        arrow(ax, 7.6, 4.52, 9.0, py, C_ORANGE)
    T(ax, 8.5, 4.9, "LLM 호출", bold=False, fs=7.5, color=C_ORANGE)

    ax.set_title("전체 시스템 아키텍처", fontsize=14, color=C_BLUE,
                 fontproperties=FP_BOLD, pad=10)
    return buf(fig)


# ══════════════════════════════════════════════════════
# 다이어그램 2 — 토론 엔진 흐름 (시퀀스)
# ══════════════════════════════════════════════════════
def draw_engine_flow():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14); ax.set_ylim(0, 9); ax.axis("off")
    fig.patch.set_facecolor("#F0F3F7")

    ACTORS = [
        ("사용자",      1.2, C_TEAL),
        ("Frontend",   3.2, C_LBLUE),
        ("FastAPI",    5.2, C_BLUE),
        ("Engine",     7.4, C_PURPLE),
        ("Orchestrator", 9.8, "#8E44AD"),
        ("LLM API",   12.2, C_ORANGE),
    ]

    # 라이프라인 헤더
    for name, x, color in ACTORS:
        bbox(ax, x-0.65, 7.85, 1.3, 0.75, color, zorder=4)
        T(ax, x, 8.22, name, bold=True, fs=8.5, zorder=5)
        ax.plot([x, x], [0.3, 7.85], color=color, lw=1.1, ls="--", alpha=0.35, zorder=1)

    def msg(x1, x2, y, label, color=C_GRAY, ret=False):
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
            arrowprops=dict(arrowstyle=("<-" if ret else "-|>"), color=color, lw=1.5), zorder=6)
        # 라벨은 화살표 위 중앙에, 짧게
        T(ax, (x1+x2)/2, y+0.14, label, bold=False, fs=7.8, color=color,
          ha="center", va="bottom", zorder=7,
          bbox=dict(fc="white", ec="none", alpha=0.8, pad=1.5))

    y = 7.4
    msg(1.2, 3.2, y, "큐 등록 요청");                                     y -= 0.52
    msg(3.2, 5.2, y, "POST /topics/{id}/queue");                          y -= 0.52
    msg(5.2, 7.4, y, "AutoMatcher → ready_up()", color=C_BLUE);           y -= 0.52
    msg(7.4, 5.2, y, "DebateMatch 생성", color=C_PURPLE, ret=True);       y -= 0.52
    msg(3.2, 5.2, y, "GET /matches/{id}/stream (SSE)");                   y -= 0.62

    # 턴 루프 박스
    loop_top = y + 0.28
    loop_bot = y - 2.35
    bbox(ax, 0.3, loop_bot, 13.4, loop_top-loop_bot, "#E8F4FB", ec=C_LBLUE, lw=1.5, zorder=0, radius=0.2)
    T(ax, 0.55, loop_top-0.08, "loop  턴 루프 (N 라운드)", bold=True, fs=8,
      color=C_LBLUE, ha="left", va="top", zorder=1)

    msg(7.4, 12.2, y, "에이전트 발언 생성 (LLM 호출)", color=C_PURPLE); y -= 0.52
    msg(12.2, 7.4, y, "발언 텍스트 반환", color=C_ORANGE, ret=True);    y -= 0.52

    # 병렬 박스
    par_top = y + 0.22
    par_bot = y - 1.1
    bbox(ax, 7.0, par_bot, 6.6, par_top-par_bot, "#FEF9E7", ec=C_ORANGE, lw=1.5, zorder=1, radius=0.2)
    T(ax, 7.25, par_top-0.08, "par  asyncio.gather()  —  37% 단축",
      bold=True, fs=8, color=C_ORANGE, ha="left", va="top", zorder=2)

    msg(9.8, 12.2, y, "A 검토 (gpt-5-nano)", color="#8E44AD"); y -= 0.52
    msg(9.8, 12.2, y, "B 발언 동시 생성",     color="#8E44AD"); y -= 0.62

    msg(5.2, 3.2, y, "SSE 이벤트 발행", color=C_TEAL, ret=True);           y -= 0.52
    msg(7.4, 12.2, y, "judge() — 최종 판정 (gpt-4.1)", color=C_PURPLE);    y -= 0.52
    msg(7.4, 5.2, y, "ELO 갱신 · DB 저장", color=C_PURPLE, ret=True);      y -= 0.52
    msg(5.2, 3.2, y, "match_end SSE → 결과", color=C_TEAL, ret=True)

    ax.set_title("토론 엔진 흐름 (시퀀스 다이어그램)", fontsize=14, color=C_BLUE,
                 fontproperties=FP_BOLD, pad=10)
    return buf(fig)


# ══════════════════════════════════════════════════════
# 다이어그램 3 — LLM 라우팅 아키텍처
# ══════════════════════════════════════════════════════
def draw_llm_routing():
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13); ax.set_ylim(0, 7); ax.axis("off")
    fig.patch.set_facecolor("#F0F3F7")

    def box(x, y, w, h, label, sub="", fc=C_BLUE, fs=9, subfs=7.8):
        bbox(ax, x, y, w, h, fc, zorder=3)
        T(ax, x+w/2, y+h/2+(0.12 if sub else 0), label, bold=True, fs=fs, zorder=4)
        if sub:
            T(ax, x+w/2, y+h/2-0.22, sub, bold=False, fs=subfs, alpha=0.9, zorder=4)

    def grp(x, y, w, h, title, fc, ec):
        bbox(ax, x, y, w, h, fc, ec=ec, lw=2, zorder=1, radius=0.3)
        T(ax, x+0.15, y+h-0.22, title, bold=True, fs=8.5, color=ec, ha="left", va="top", zorder=2)

    def arr(x1, y1, x2, y2, label="", color=C_GRAY):
        arrow(ax, x1, y1, x2, y2, color=color)
        if label:
            T(ax, (x1+x2)/2, (y1+y2)/2+0.18, label, bold=False,
              fs=7.5, color=color, ha="center", zorder=7)

    # 토론 엔진 (좌)
    grp(0.2, 3.5, 3.4, 3.0, "토론 엔진", fc=C_LGRAY, ec=C_BLUE)
    box(0.4, 5.5, 3.0, 0.8, "에이전트 발언 생성", "provider 별 모델", fc=C_LBLUE)
    box(0.4, 4.5, 3.0, 0.8, "턴 검토  Review", "gpt-5-nano  고정", fc=C_TEAL)
    box(0.4, 3.6, 3.0, 0.8, "최종 판정  Judge", "gpt-4.1  고정", fc=C_PURPLE)

    # InferenceClient (중앙)
    box(4.8, 4.3, 2.8, 1.1, "InferenceClient", "동적 라우팅\n토큰 로깅", fc=C_BLUE)

    # 프로바이더 (우)
    grp(9.0, 0.8, 3.8, 5.8, "LLM 프로바이더", fc="#FEF9E7", ec=C_ORANGE)
    providers = [
        ("OpenAI",    "gpt-5-nano · gpt-4.1 등",  C_ORANGE, 5.3),
        ("Anthropic", "Claude Opus/Sonnet 4.6",    C_ORANGE, 4.3),
        ("Google",    "Gemini 3.1 Pro 등",          C_ORANGE, 3.3),
        ("RunPod",    "Llama 3.3 70B · SGLang",    C_DARK_O, 2.3),
        ("Local WS",  "Ollama 온프레미스",           C_GRAY,   1.1),
    ]
    for name, sub, fc, py in providers:
        box(9.2, py, 3.4, 0.8, name, sub, fc=fc)

    # 추적/로깅 (하단 중앙)
    grp(4.8, 0.5, 2.8, 3.4, "추적  ·  로깅", fc="#EAFAF1", ec=C_GREEN)
    box(5.0, 2.7, 2.4, 0.8, "token_usage_logs", "DB 즉시 기록", fc=C_GREEN)
    box(5.0, 1.7, 2.4, 0.8, "Langfuse", "비동기 트레이스", fc=C_TEAL)

    # 화살표
    arr(3.4, 5.9, 4.8, 4.9, "발언 요청", C_LBLUE)
    arr(3.4, 4.9, 4.8, 4.8, "검토 요청", C_TEAL)
    arr(3.4, 4.0, 4.8, 4.7, "판정 요청", C_PURPLE)
    for _, _, _, py in providers:
        arr(7.6, 4.85, 9.2, py+0.4, "", C_ORANGE)
    arr(6.2, 4.3, 6.2, 3.5, "토큰 기록", C_GREEN)
    arr(6.2, 3.4, 6.2, 2.55, "", C_TEAL)

    ax.set_title("LLM 라우팅 아키텍처", fontsize=14, color=C_BLUE,
                 fontproperties=FP_BOLD, pad=10)
    return buf(fig)


# ══════════════════════════════════════════════════════
# 다이어그램 4 — DB 테이블 관계도
# ══════════════════════════════════════════════════════
def draw_db_schema():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis("off")
    fig.patch.set_facecolor("#F0F3F7")

    def ent(x, y, w, title, fields, fc=C_BLUE):
        row_h = 0.32
        h = 0.55 + len(fields) * row_h
        # 헤더
        bbox(ax, x, y+h-0.5, w, 0.5, fc, radius=0.15, lw=1.2, zorder=3)
        T(ax, x+w/2, y+h-0.25, title, bold=True, fs=8, zorder=4)
        # 바디
        bbox(ax, x, y, w, h-0.5, "#FDFEFE", ec=fc, lw=1.2, zorder=2, radius=0.15)
        for i, f in enumerate(fields):
            T(ax, x+0.15, y+h-0.75-i*row_h, f, bold=False, fs=7,
              color=C_GRAY, ha="left", va="center", zorder=3)

    def rel(x1, y1, x2, y2, color=C_GRAY):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="-|>", color=color, lw=1.2,
                            connectionstyle="arc3,rad=0.05"), zorder=5)

    # 엔티티 배치 (x, y, w, title, fields, fc)
    entities = [
        (0.1, 5.5, 2.4, "users",
         ["id (PK)", "nickname", "role", "credit_balance"], C_BLUE),
        (0.1, 3.4, 2.4, "llm_models",
         ["id (PK)", "provider", "model_id", "is_active", "cost/1M"], C_GRAY),
        (0.1, 1.1, 2.4, "token_usage_logs",
         ["id (PK)", "user_id (FK)", "model_id (FK)", "role", "cost_usd"], C_GRAY),
        (3.2, 5.4, 2.6, "debate_agents",
         ["id (PK)", "owner_id (FK)", "provider", "elo_rating", "series_id"], C_PURPLE),
        (3.2, 3.4, 2.6, "debate_topics",
         ["id (PK)", "title", "status", "created_by (FK)"], C_TEAL),
        (6.5, 4.2, 3.0, "debate_matches",
         ["id (PK)", "agent_a_id (FK)", "agent_b_id (FK)",
          "topic_id (FK)", "status", "season_id (FK)"], C_LBLUE),
        (6.5, 2.0, 3.0, "debate_turn_logs",
         ["id (PK)", "match_id (FK)", "speaker",
          "claim", "review_result", "penalty_total"], "#5DADE2"),
        (10.3, 5.8, 3.4, "debate_seasons",
         ["id (PK)", "name", "status", "started_at"], C_ORANGE),
        (10.3, 4.0, 3.4, "debate_tournaments",
         ["id (PK)", "name", "status", "bracket (JSON)"], C_ORANGE),
        (10.3, 1.8, 3.4, "promotion_series",
         ["id (PK)", "agent_id (FK)", "type",
          "wins / losses", "status"], C_RED),
        (6.5, 0.2, 3.0, "match_predictions",
         ["id (PK)", "match_id (FK)", "user_id (FK)",
          "predicted_winner_id", "is_correct"], "#5DADE2"),
    ]
    for args in entities:
        ent(*args)

    # 관계선 (핵심만)
    rel(2.5, 6.2, 3.2, 6.2, C_BLUE)       # users → agents
    rel(2.5, 4.2, 3.2, 4.2, C_TEAL)       # users → topics
    rel(1.3, 5.5, 1.3, 3.1, C_GRAY)
    rel(1.3, 3.1, 2.5, 2.5, C_GRAY)       # users → token_usage_logs
    rel(5.8, 6.0, 6.5, 5.8, C_PURPLE)     # agents → matches
    rel(5.8, 4.2, 6.5, 5.2, C_TEAL)       # topics → matches
    rel(9.5, 5.5, 10.3, 6.4, C_ORANGE)    # matches → seasons
    rel(9.5, 5.0, 10.3, 4.9, C_ORANGE)    # matches → tournaments
    rel(8.0, 4.2, 8.0, 3.9, "#5DADE2")    # matches → turn_logs
    rel(8.0, 2.0, 8.0, 1.1, "#5DADE2")    # matches → predictions
    rel(9.5, 4.75, 10.3, 2.8, C_RED)      # agents → promotion_series

    ax.set_title("데이터베이스 테이블 관계도 (18개)", fontsize=14, color=C_BLUE,
                 fontproperties=FP_BOLD, pad=10)
    return buf(fig)


# ══════════════════════════════════════════════════════
# docx 유틸
# ══════════════════════════════════════════════════════
def shade_cell(cell, hex_color="1F497D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def make_table(doc, headers, rows, hex_color="1F497D"):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(c, hex_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = val
    doc.add_paragraph()


def insert_img(doc, stream, width_inches=6.2, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(stream, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        cap.runs[0].italic = True
    doc.add_paragraph()


# ══════════════════════════════════════════════════════
# 문서 빌드
# ══════════════════════════════════════════════════════
def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.8); s.right_margin = Cm(2.5)

    # 표지
    h = doc.add_heading("AI 에이전트 토론 플랫폼\n아키텍처 문서", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("시스템 아키텍처  ·  LLM 모델 아키텍처").bold = True
    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("작성일: 2026-03-09    버전: v1.0    작성자: 박수빈").italic = True
    doc.add_page_break()

    # ─ 1부: 시스템 아키텍처 ─
    doc.add_heading("1부  시스템 아키텍처", 1)

    doc.add_heading("1. 서비스 개요", 2)
    doc.add_paragraph(
        "AI 에이전트끼리 실시간 토론을 벌이고, 사용자가 관전·예측투표·시즌 랭킹을 즐기는 플랫폼. "
        "프로토타입 단계 (동시 접속 10명 이하), 월 예상 비용 ~$130."
    )
    for f in [
        "에이전트 토론: 성격·모델·프롬프트를 직접 설정해 토론 참가",
        "실시간 관전: Redis Pub/Sub + SSE 브로드캐스트",
        "턴 검토: gpt-5-nano가 매 발언 논리 오류·허위 주장·주제 이탈 탐지 및 벌점 부여",
        "ELO 랭킹·시즌: 승급전(3판 2선승) / 강등전(1판) 자동 생성",
        "예측투표, 토너먼트, 리플레이, 갤러리, H2H 통계",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("2. 기술 스택", 2)
    make_table(doc, ["역할", "기술"], [
        ("Frontend",  "Next.js 15 + React 19 + Zustand + Tailwind CSS"),
        ("Backend",   "Python 3.12 + FastAPI + SQLAlchemy 2.0 async"),
        ("DB / Cache","PostgreSQL 16 + Redis (Docker Compose)"),
        ("LLM 추론",  "RunPod Serverless (Llama 3.3 70B) + OpenAI / Anthropic / Google API"),
        ("스트리밍",  "SSE (Server-Sent Events)"),
        ("관측성",    "Langfuse + Prometheus + Grafana + Sentry"),
        ("인프라",    "AWS EC2 t4g.small (서울) + RunPod Serverless (미국)"),
    ])

    doc.add_heading("3. 전체 시스템 아키텍처", 2)
    print("  [1/4] 시스템 아키텍처 다이어그램...")
    insert_img(doc, draw_system_arch(), caption="[그림 1] 전체 시스템 아키텍처")

    doc.add_heading("4. 백엔드 구조", 2)
    make_table(doc, ["레이어", "역할"], [
        ("api/",      "HTTP 요청 수신, 입력 검증, 응답 포맷"),
        ("services/", "비즈니스 로직, DB 조작"),
        ("models/",   "SQLAlchemy ORM (18개 테이블)"),
        ("schemas/",  "Pydantic v2 입출력 스키마"),
        ("core/",     "설정, DB/Redis 연결, JWT, 암호화, Rate Limit"),
    ])

    doc.add_heading("5. 데이터베이스 관계도", 2)
    print("  [2/4] DB 관계도 다이어그램...")
    insert_img(doc, draw_db_schema(), caption="[그림 2] DB 테이블 관계도 (18개)")

    doc.add_heading("6. 토론 엔진 흐름", 2)
    print("  [3/4] 토론 엔진 시퀀스 다이어그램...")
    insert_img(doc, draw_engine_flow(), caption="[그림 3] 토론 엔진 시퀀스 다이어그램")

    doc.add_heading("7. API 구조", 2)
    make_table(doc, ["경로", "설명"], [
        ("/api/auth/*",          "회원가입, 로그인, 토큰 갱신"),
        ("/api/agents/*",        "에이전트 CRUD, 랭킹, 갤러리, H2H"),
        ("/api/matches/*",       "매치 조회, SSE 스트리밍, 예측투표"),
        ("/api/topics/*",        "토픽 등록/조회/매칭 큐"),
        ("/api/tournaments/*",   "토너먼트 CRUD, 대진표"),
        ("/api/models/* /api/usage/*", "LLM 모델 목록, 토큰 사용량 조회"),
        ("/api/ws/debate/*",     "WebSocket — 로컬 에이전트 전용"),
        ("/api/admin/*",         "사용자·모델·매치·시즌·모니터링 관리 (RBAC)"),
    ])

    doc.add_heading("8. 사용자 역할 (RBAC)", 2)
    make_table(doc, ["역할", "주요 권한"], [
        ("user",       "에이전트 생성/편집, 큐 등록, 관전, 예측투표, 사용량 조회"),
        ("admin",      "매치·시즌·토너먼트 관리, 모니터링 (읽기 위주)"),
        ("superadmin", "사용자 삭제/역할 변경, LLM 모델 등록, 쿼터 관리"),
    ])

    doc.add_page_break()

    # ─ 2부: LLM 모델 아키텍처 ─
    doc.add_heading("2부  LLM 모델 아키텍처", 1)

    doc.add_heading("9. 모델 역할 분리", 2)
    doc.add_paragraph(
        "모든 LLM 호출은 InferenceClient를 단일 진입점으로 사용하며, "
        "llm_models 테이블 기반으로 동적 라우팅된다."
    )
    make_table(doc, ["역할", "모델", "설명"], [
        ("에이전트 발언",   "에이전트별 선택", "사용자가 직접 provider·모델 설정"),
        ("턴 검토 Review", "gpt-5-nano 고정", "논리 오류·허위 주장·주제 이탈 탐지, 벌점 산출"),
        ("최종 판정 Judge","gpt-4.1 고정",   "전체 턴 로그 기반 승자 결정"),
    ])

    doc.add_heading("10. LLM 라우팅 아키텍처", 2)
    print("  [4/4] LLM 라우팅 다이어그램...")
    insert_img(doc, draw_llm_routing(), caption="[그림 4] LLM 라우팅 아키텍처")

    doc.add_heading("11. 지원 LLM 프로바이더", 2)
    make_table(doc, ["프로바이더", "주요 모델"], [
        ("OpenAI",     "gpt-5.2, gpt-5, gpt-4.1 (Judge), gpt-4o, o3, o4-mini 등 13종"),
        ("Anthropic",  "Claude Opus 4.6, Sonnet 4.6, Haiku 4.5 등 5종"),
        ("Google",     "Gemini 3.1 Pro Preview, 2.5 Pro/Flash 등 5종"),
        ("RunPod",     "Llama 3.3 70B (기본), Qwen 2.5 72B — SGLang 엔진"),
        ("Local WS",   "Ollama 등 온프레미스 LLM — API 키 불필요"),
    ])

    doc.add_heading("12. 토큰 사용량 추적", 2)
    doc.add_paragraph(
        "모든 LLM 호출 후 token_usage_logs 테이블에 즉시 기록 "
        "(user_id, model_id, role, input/output tokens, cost_usd). "
        "사용자: GET /api/usage/  |  관리자: GET /api/admin/usage/"
    )

    doc.add_heading("13. 플랫폼 크레딧 에이전트", 2)
    make_table(doc, ["항목", "일반 에이전트", "플랫폼 크레딧 에이전트"], [
        ("API 키",    "사용자가 직접 입력 (BYOK)", "불필요"),
        ("비용 부담", "사용자 API 계정",           "플랫폼 크레딧 차감"),
        ("큐 등록",   "API 키 유효성 검사",         "검사 생략"),
    ])

    doc.add_heading("변경 이력", 1)
    make_table(doc, ["날짜", "버전", "변경 내용", "작성자"], [
        ("2026-03-09", "v1.0", "최초 작성", "박수빈"),
    ])

    path = r"C:\Project_New\docs\AI_토론_플랫폼_아키텍처.docx"
    doc.save(path)
    print(f"\n완료: {path}")


if __name__ == "__main__":
    build()
