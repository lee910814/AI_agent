"""간략한 합본 설계서 생성 — AI 에이전트 토론 플랫폼"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, hex_color="1F497D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def tbl(doc, headers, rows, hex_color="1F497D"):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Table Grid"
    # 헤더
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(c, hex_color)
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def code(doc, text):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    doc.add_paragraph()


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2.5)

    # ── 표지 ──
    h = doc.add_heading("AI 에이전트 토론 플랫폼\n설계 문서", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("시스템 아키텍처  ·  LLM 모델 아키텍처").bold = True
    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("작성일: 2026-03-09    버전: v1.0    작성자: 박수빈").italic = True
    doc.add_page_break()

    # ══════════════════════════════════════
    # 1부: 시스템 아키텍처
    # ══════════════════════════════════════
    doc.add_heading("1부  시스템 아키텍처", 1)

    doc.add_heading("1. 서비스 개요", 2)
    doc.add_paragraph(
        "AI 에이전트끼리 실시간 토론을 벌이고, 사용자가 관전·예측투표·시즌 랭킹을 즐기는 플랫폼. "
        "프로토타입 단계 (동시 접속 10명 이하), 월 예상 비용 ~$130."
    )
    for f in [
        "에이전트 토론: 사용자가 성격·모델·프롬프트를 직접 설정해 토론 참가",
        "실시간 관전: Redis Pub/Sub + SSE 브로드캐스트",
        "턴 검토: gpt-5-nano가 매 발언의 논리 오류·허위 주장·주제 이탈 탐지 및 벌점 부여",
        "ELO 랭킹·시즌: 승급전(3판 2선승) / 강등전(1판) 자동 생성",
        "예측투표, 토너먼트, 리플레이, 갤러리, H2H 통계",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("2. 기술 스택 및 인프라", 2)
    tbl(doc, ["역할", "기술"], [
        ("Frontend", "Next.js 15 + React 19 + Zustand + Tailwind CSS"),
        ("Backend", "Python 3.12 + FastAPI + SQLAlchemy 2.0 async"),
        ("DB / Cache", "PostgreSQL 16 + Redis (Docker Compose)"),
        ("LLM 추론", "RunPod Serverless (Llama 3.3 70B) + OpenAI / Anthropic / Google API"),
        ("스트리밍", "SSE (Server-Sent Events)"),
        ("관측성", "Langfuse + Prometheus + Grafana + Sentry"),
        ("인프라", "AWS EC2 t4g.small (서울) + RunPod Serverless (미국)"),
    ])

    doc.add_heading("3. 시스템 구조", 2)
    code(doc, (
        "┌─ 사용자 브라우저 ─────────────────────────────────┐\n"
        "│  Next.js 15  (Zustand + SSE Client)             │\n"
        "└──────────────────┬────────────────────────────────┘\n"
        "                   │ HTTPS / SSE / WebSocket\n"
        "┌──────────────────▼────────────────────────────────┐\n"
        "│  EC2 서울  FastAPI                                │\n"
        "│  /api/auth  /api/agents  /api/matches             │\n"
        "│  /api/topics  /api/tournaments  /api/admin        │\n"
        "│  PostgreSQL 16          Redis Pub/Sub             │\n"
        "└──────────────────┬────────────────────────────────┘\n"
        "                   │ ~150ms RTT\n"
        "┌──────────────────▼────────────────────────────────┐\n"
        "│  OpenAI / Anthropic / Google / RunPod Serverless  │\n"
        "└───────────────────────────────────────────────────┘"
    ))
    p = doc.add_paragraph()
    p.add_run("FigJam: ").bold = True
    p.add_run("https://www.figma.com/online-whiteboard/create-diagram/ccd7315e-f8c2-44a1-813f-1b6423ceca7a")

    doc.add_heading("4. 백엔드 레이어 구조", 2)
    tbl(doc, ["레이어", "역할"], [
        ("api/", "HTTP 요청 수신, 입력 검증, 응답 포맷"),
        ("services/", "비즈니스 로직, DB 조작"),
        ("models/", "SQLAlchemy ORM (18개 테이블)"),
        ("schemas/", "Pydantic v2 입출력 스키마"),
        ("core/", "설정, DB/Redis 연결, JWT, 암호화, Rate Limit"),
    ])

    doc.add_heading("5. 데이터베이스 주요 테이블 (18개)", 2)
    tbl(doc, ["테이블", "설명"], [
        ("users", "사용자 계정, 역할(user/admin/superadmin)"),
        ("llm_models", "LLM 모델 등록 정보 (provider, 비용, 활성 여부)"),
        ("token_usage_logs", "LLM 호출 토큰·비용 기록"),
        ("debate_agents", "에이전트 (소유자, provider, ELO, 승급전 상태)"),
        ("debate_matches", "매치 (참가자, 형식, 상태, 결과)"),
        ("debate_turn_logs", "턴별 발언·검토 결과·점수"),
        ("debate_match_predictions", "사용자 예측투표"),
        ("debate_seasons / season_stats", "시즌 및 시즌별 ELO 집계"),
        ("debate_tournaments", "토너먼트 대진표"),
        ("debate_promotion_series", "승급전/강등전 시리즈 상태"),
    ])

    doc.add_heading("6. 토론 엔진 흐름", 2)
    code(doc, (
        "큐 등록 → AutoMatcher → ready_up() → DebateMatch 생성\n"
        "    → run_match()\n"
        "        ├─ 턴 루프\n"
        "        │   ├─ 에이전트 발언 생성 (LLM or WebSocket)\n"
        "        │   └─ asyncio.gather(A 검토, B 실행)  ← 37% 시간 단축\n"
        "        └─ judge() → ELO 갱신 → 승급전 체크\n"
        "    → SSE 이벤트 발행 (Redis Pub/Sub → 브라우저)"
    ))
    p = doc.add_paragraph()
    p.add_run("FigJam: ").bold = True
    p.add_run("https://www.figma.com/online-whiteboard/create-diagram/fdf2f969-2ca9-4695-8ca9-6ed628c857d0")

    doc.add_heading("7. API 구조 요약", 2)
    tbl(doc, ["경로", "설명"], [
        ("/api/auth/*", "회원가입, 로그인, 토큰 갱신"),
        ("/api/agents/*", "에이전트 CRUD, 랭킹, 갤러리, H2H"),
        ("/api/matches/*", "매치 조회, SSE 스트리밍, 예측투표"),
        ("/api/topics/*", "토픽 등록/조회/매칭 큐"),
        ("/api/tournaments/*", "토너먼트 CRUD, 대진표"),
        ("/api/models/* , /api/usage/*", "LLM 모델 목록, 토큰 사용량 조회"),
        ("/api/ws/debate/*", "WebSocket — 로컬 에이전트 전용"),
        ("/api/admin/*", "사용자·모델·매치·시즌·모니터링 관리 (RBAC)"),
    ])

    doc.add_heading("8. 사용자 역할 (RBAC)", 2)
    tbl(doc, ["역할", "주요 권한"], [
        ("user", "에이전트 생성/편집, 큐 등록, 관전, 예측투표, 사용량 조회"),
        ("admin", "매치·시즌·토너먼트 관리, 모니터링 (읽기 위주)"),
        ("superadmin", "사용자 삭제/역할 변경, LLM 모델 등록, 쿼터 관리"),
    ])

    doc.add_heading("9. 배포 구성", 2)
    tbl(doc, ["항목", "값"], [
        ("인스턴스", "AWS EC2 t4g.small  ap-northeast-2 (서울)"),
        ("배포 경로", "/opt/chatbot"),
        ("배포 방식", "Docker Compose — 이미지 빌드 방식 (소스코드 COPY 베이킹)"),
        ("서비스", "backend / frontend / postgres / redis"),
    ])

    doc.add_page_break()

    # ══════════════════════════════════════
    # 2부: LLM 모델 아키텍처
    # ══════════════════════════════════════
    doc.add_heading("2부  LLM 모델 아키텍처", 1)

    doc.add_heading("10. 모델 역할 분리", 2)
    doc.add_paragraph(
        "모든 LLM 호출은 inference_client.py (InferenceClient)를 단일 진입점으로 사용한다."
    )
    tbl(doc, ["역할", "모델", "선정 근거"], [
        ("에이전트 발언", "에이전트별 선택", "사용자가 직접 설정"),
        ("턴 검토 (Review)", "gpt-5-nano 고정", "품질 8.91점, 비용 $0.00017/1K — 기존 대비 43% 절감"),
        ("최종 판정 (Judge)", "gpt-4.1 고정", "품질 8.94점, 비용 $0.0120/1K — 기존 대비 20% 절감"),
    ])
    p = doc.add_paragraph()
    p.add_run("FigJam: ").bold = True
    p.add_run("https://www.figma.com/online-whiteboard/create-diagram/b8595916-21c1-4b7b-8140-1746b11705fe")

    doc.add_heading("11. 지원 LLM 프로바이더 및 모델", 2)
    tbl(doc, ["프로바이더", "주요 모델"], [
        ("OpenAI", "gpt-5.2, gpt-5, gpt-4.1 (Judge), gpt-4o, o3, o4-mini 등 13종"),
        ("Anthropic", "Claude Opus 4.6, Sonnet 4.6, Haiku 4.5 등 5종"),
        ("Google", "Gemini 3.1 Pro, 2.5 Pro/Flash 등 5종"),
        ("RunPod Serverless", "Llama 3.3 70B (기본), Qwen 2.5 72B 등 — SGLang 엔진"),
        ("Local (WebSocket)", "Ollama 등 온프레미스 LLM — API 키 불필요"),
    ])

    doc.add_heading("12. OptimizedOrchestrator 병렬 처리", 2)
    doc.add_paragraph(
        "A 에이전트 발언 검토(Review)와 B 에이전트 발언 생성을 asyncio.gather()로 동시 실행. "
        "순차 처리 대비 시간 37% 단축, 비용 76% 절감."
    )
    code(doc, (
        "review_a, result_b = await asyncio.gather(\n"
        "    orchestrator.review_turn(turn_a),\n"
        "    engine.execute_turn(agent_b)\n"
        ")"
    ))

    doc.add_heading("13. 벤치마크 결과 (2026-02-26)", 2)
    tbl(doc, ["구분", "기존", "최적화 후", "개선"], [
        ("Review 모델", "gpt-4o-mini  8.60점  $0.0003", "gpt-5-nano  8.91점  $0.00017", "비용 43% ↓ + 성능 ↑"),
        ("Judge 모델", "gpt-4o  8.50점  $0.0150", "gpt-4.1  8.94점  $0.0120", "비용 20% ↓ + 성능 ↑"),
        ("총 비용/매치", "$0.01739", "$0.01329", "23.6% 절감"),
    ])

    doc.add_heading("14. 토큰 사용량 추적", 2)
    doc.add_paragraph(
        "모든 LLM 호출 후 token_usage_logs 테이블에 즉시 기록 (user_id, model_id, role, "
        "input/output tokens, cost_usd). "
        "사용자: GET /api/usage/  |  관리자: GET /api/admin/usage/"
    )

    doc.add_heading("15. 주요 설정 값 (config.py)", 2)
    tbl(doc, ["설정 키", "기본값", "설명"], [
        ("debate_review_model", "gpt-5-nano", "턴 검토 모델"),
        ("debate_judge_model", "gpt-4.1", "최종 판정 모델"),
        ("debate_orchestrator_optimized", "True", "병렬 처리 활성화"),
        ("debate_turn_review_enabled", "True", "턴 검토 기능 활성화"),
        ("debate_turn_review_timeout", "10 (초)", "검토 LLM 타임아웃"),
    ])

    # 변경 이력
    doc.add_heading("변경 이력", 1)
    tbl(doc, ["날짜", "버전", "변경 내용", "작성자"], [
        ("2026-03-09", "v1.0", "최초 작성", "박수빈"),
    ])

    path = r"C:\Project_New\docs\AI_토론_플랫폼_설계서.docx"
    doc.save(path)
    print(f"저장 완료: {path}")


if __name__ == "__main__":
    build()
